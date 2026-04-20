from __future__ import annotations

import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from replacedocx.engine_lib.common import BADGE_TAG, SECTION_TAG  # noqa: E402
from replacedocx.cli import build_parser  # noqa: E402
from replacedocx.engine_lib.docx_utils import replace_question_prefix_marker_in_paragraph  # noqa: E402
from replacedocx.engine_lib.pipeline import processar_docx  # noqa: E402


class ReplaceQuestionPrefixMarkerTests(unittest.TestCase):
    def _render_document_xml(self, with_math: bool) -> str:
        with TemporaryDirectory() as td:
            doc_path = Path(td) / "sample.docx"

            doc = Document()
            p = doc.add_paragraph()
            p.add_run("1. (MÉDIA) Energia ")
            if with_math:
                omml = parse_xml(
                    rf"""
                    <m:oMath {nsdecls('m')}>
                      <m:r><m:t>E=mc2</m:t></m:r>
                    </m:oMath>
                    """
                )
                p._element.append(omml)
            p.add_run(" final")
            doc.save(doc_path)

            reopened = Document(doc_path)
            replaced = replace_question_prefix_marker_in_paragraph(
                reopened.paragraphs[0],
                {"(MÉDIA)": "Assets/areas/fisica/capsulas/media.png"},
                1.3,
                "Arial",
                11,
            )
            self.assertTrue(replaced)

            reopened.save(doc_path)
            return ZipFile(doc_path).read("word/document.xml").decode("utf-8")

    def test_replaces_prefix_marker_in_plain_paragraph(self) -> None:
        xml = self._render_document_xml(with_math=False)

        self.assertIn(BADGE_TAG, xml)
        self.assertNotIn("(MÉDIA)", xml)
        self.assertIn("Energia", xml)

    def test_preserves_office_math_when_replacing_prefix_marker(self) -> None:
        xml = self._render_document_xml(with_math=True)

        self.assertIn(BADGE_TAG, xml)
        self.assertIn("m:oMath", xml)
        self.assertIn("Energia", xml)


class AreaPipelineRegressionTests(unittest.TestCase):
    def _render_processed_document_xml(self, area: str) -> str:
        with TemporaryDirectory() as td:
            td_path = Path(td)
            input_docx = td_path / f"{area}.docx"
            output_docx = td_path / f"{area}_ok.docx"

            doc = Document()
            doc.add_paragraph("EXERCÍCIOS PROPOSTOS")
            p = doc.add_paragraph()
            p.add_run("1. (MÉDIA) Energia ")
            p._element.append(
                parse_xml(
                    rf"""
                    <m:oMath {nsdecls('m')}>
                      <m:r><m:t>E=mc2</m:t></m:r>
                    </m:oMath>
                    """
                )
            )
            p.add_run(" final")
            doc.save(input_docx)

            processar_docx(
                input_docx,
                output_docx,
                {
                    "area_conhecimento": area,
                    "finalize_word": False,
                    "insert_section_banners": False,
                    "insert_question_tables": False,
                    "append_difficulty_report": False,
                },
            )

            return ZipFile(output_docx).read("word/document.xml").decode("utf-8")

    def test_quimica_pipeline_preserves_formula_when_replacing_badge(self) -> None:
        xml = self._render_processed_document_xml("quimica")

        self.assertIn(BADGE_TAG, xml)
        self.assertIn("m:oMath", xml)
        self.assertNotIn("(MÉDIA)", xml)

    def test_matematica_pipeline_preserves_formula_when_replacing_badge(self) -> None:
        xml = self._render_processed_document_xml("matematica")

        self.assertIn(BADGE_TAG, xml)
        self.assertIn("m:oMath", xml)
        self.assertNotIn("(MÉDIA)", xml)

    def _render_section_banner_document_xml(self, area: str) -> str:
        with TemporaryDirectory() as td:
            td_path = Path(td)
            input_docx = td_path / f"{area}.docx"
            output_docx = td_path / f"{area}_ok.docx"

            doc = Document()
            p = doc.add_paragraph()
            p.add_run("EXERCÍCIOS PROPOSTOS: 1. A energia ")
            p._element.append(
                parse_xml(
                    rf"""
                    <m:oMath {nsdecls('m')}>
                      <m:r><m:t>E=mc2</m:t></m:r>
                    </m:oMath>
                    """
                )
            )
            p.add_run(" permanece.")
            doc.save(input_docx)

            processar_docx(
                input_docx,
                output_docx,
                {
                    "area_conhecimento": area,
                    "finalize_word": False,
                    "insert_question_tables": False,
                    "append_difficulty_report": False,
                },
            )

            return ZipFile(output_docx).read("word/document.xml").decode("utf-8")

    def test_fisica_pipeline_preserves_enunciado_when_section_banner_is_inserted(self) -> None:
        xml = self._render_section_banner_document_xml("fisica")

        self.assertIn(SECTION_TAG, xml)
        self.assertIn("m:oMath", xml)
        self.assertIn("A energia", xml)
        self.assertIn("permanece.", xml)
        self.assertNotIn("EXERCÍCIOS PROPOSTOS: 1. A energia", xml)

    def test_quimica_pipeline_preserves_enunciado_when_section_banner_is_inserted(self) -> None:
        xml = self._render_section_banner_document_xml("quimica")

        self.assertIn(SECTION_TAG, xml)
        self.assertIn("m:oMath", xml)
        self.assertIn("A energia", xml)
        self.assertIn("permanece.", xml)
        self.assertNotIn("EXERCÍCIOS PROPOSTOS: 1. A energia", xml)

    def test_matematica_pipeline_preserves_enunciado_when_section_banner_is_inserted(self) -> None:
        xml = self._render_section_banner_document_xml("matematica")

        self.assertIn(SECTION_TAG, xml)
        self.assertIn("m:oMath", xml)
        self.assertIn("A energia", xml)
        self.assertIn("permanece.", xml)
        self.assertNotIn("EXERCÍCIOS PROPOSTOS: 1. A energia", xml)

    def test_preserve_paragraphs_keeps_answer_lines(self) -> None:
        with TemporaryDirectory() as td:
            td_path = Path(td)
            input_docx = td_path / "fisica_preserve.docx"
            output_docx = td_path / "fisica_preserve_ok.docx"

            doc = Document()
            doc.add_paragraph("EXERCÍCIOS PROPOSTOS")
            doc.add_paragraph("1. Texto da questão.")
            doc.add_paragraph("Resposta: A")
            doc.save(input_docx)

            processar_docx(
                input_docx,
                output_docx,
                {
                    "area_conhecimento": "fisica",
                    "finalize_word": False,
                    "insert_section_banners": False,
                    "insert_question_tables": False,
                    "append_difficulty_report": False,
                    "preserve_paragraphs": True,
                },
            )

            xml = ZipFile(output_docx).read("word/document.xml").decode("utf-8")
            self.assertIn("Resposta: A", xml)


class CliRegressionTests(unittest.TestCase):
    def test_process_parser_accepts_preserve_paragraphs_flag(self) -> None:
        args = build_parser().parse_args(
            ["process", "entrada.docx", "--area", "fisica", "--preserve-paragraphs"]
        )

        self.assertTrue(args.preserve_paragraphs)


if __name__ == "__main__":
    unittest.main()
