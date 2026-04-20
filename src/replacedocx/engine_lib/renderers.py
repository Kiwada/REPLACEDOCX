from __future__ import annotations

import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .common import SECTION_TAG, add_picture_resilient, elog, normalize_text_key, resolve_path, safe_tag_suffix
from .docx_utils import (
    iter_paragraphs,
    set_cell_fill,
    set_cell_no_wrap,
    set_cell_width,
    set_table_width,
    style_cell_text,
    trim_paragraph_leading_text,
)
from .report import difficulty_label, match_section_for_report, section_aliases_for_report


def _matches_section_title(norm_txt: str, norm_title: str) -> bool:
    if norm_txt == norm_title:
        return True
    if norm_txt.startswith(norm_title + ":"):
        return True
    if (
        norm_txt.startswith(norm_title + " -")
        or norm_txt.startswith(norm_title + " –")
        or norm_txt.startswith(norm_title + " —")
    ):
        return True
    return False


def _normalized_aliases_for_section(canonical_title: str | None, raw_title: str | None) -> list[str]:
    aliases: list[str] = []
    if canonical_title:
        aliases.append(canonical_title)
        aliases.extend(section_aliases_for_report().get(canonical_title, []))
    if raw_title:
        aliases.append(raw_title)

    seen: set[str] = set()
    normalized: list[str] = []
    for alias in aliases:
        norm_alias = normalize_text_key(alias)
        if not norm_alias or norm_alias in seen:
            continue
        seen.add(norm_alias)
        normalized.append(norm_alias)
    normalized.sort(key=len, reverse=True)
    return normalized


def _section_paragraph_has_trailing_content(norm_txt: str, normalized_aliases: list[str]) -> bool:
    for norm_alias in normalized_aliases:
        if not norm_txt.startswith(norm_alias):
            continue
        remainder = norm_txt[len(norm_alias) :]
        cleaned = re.sub(r"[\s:–—\-]+", "", remainder)
        return bool(cleaned)
    return False


def _find_section_prefix_end(text: str, aliases: list[str]) -> int | None:
    seen: set[str] = set()
    candidates = sorted((alias for alias in aliases if alias), key=len, reverse=True)
    for alias in candidates:
        if alias in seen:
            continue
        seen.add(alias)
        match = re.match(
            rf"^\s*{re.escape(alias)}(?:\s*[:\-–—]?\s*)",
            text,
            flags=re.IGNORECASE,
        )
        if match and text[match.end() :].strip():
            return match.end()
    return None


def _set_section_single_column(section) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "1")
    cols.set(qn("w:space"), "0")


def _replace_paragraph_with_section_banner(
    p,
    img_path,
    section_banner_width_cm: float,
    section_tag: str,
    alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
) -> None:
    # Substitui totalmente o título da seção pelo banner.
    p.clear()
    banner_run = p.add_run()
    pic = add_picture_resilient(banner_run, img_path, section_banner_width_cm)

    p.alignment = alignment
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    try:
        pic._inline.docPr.set("descr", section_tag)
        pic._inline.docPr.set("title", section_tag)
    except Exception:
        pass


def apply_section_banners(
    doc: Document,
    section_banners: dict[str, str],
    section_banner_width_cm: float,
) -> int:
    normalized_map: dict[str, str] = {}
    for title, img in section_banners.items():
        norm_title = normalize_text_key(title)
        if norm_title:
            normalized_map[norm_title] = img

    inserted = 0
    for p in list(iter_paragraphs(doc)):
        txt = (p.text or "").strip()
        if not txt:
            continue

        norm_txt = normalize_text_key(txt)
        if not norm_txt or len(norm_txt) > 120:
            continue

        matched_title = None
        matched_img = None
        matched_canonical = None
        canonical = match_section_for_report(norm_txt)
        if canonical:
            norm_canonical = normalize_text_key(canonical)
            if norm_canonical in normalized_map:
                matched_title = norm_canonical
                matched_img = normalized_map[norm_canonical]
                matched_canonical = canonical

        if matched_img is None:
            for norm_title, img in normalized_map.items():
                if _matches_section_title(norm_txt, norm_title):
                    matched_title = norm_title
                    matched_img = img
                    break

        if matched_img is None:
            continue

        section_tag = f"{SECTION_TAG}_{safe_tag_suffix(matched_title)}"
        normalized_aliases = _normalized_aliases_for_section(matched_canonical, matched_title)

        img_path = resolve_path(matched_img)
        if not img_path.exists():
            elog(f"Section banner missing for '{txt}': {img_path}")
            continue

        try:
            if _section_paragraph_has_trailing_content(norm_txt, normalized_aliases):
                banner_p = p.insert_paragraph_before("")
                _replace_paragraph_with_section_banner(
                    banner_p,
                    img_path=img_path,
                    section_banner_width_cm=section_banner_width_cm,
                    section_tag=section_tag,
                )

                prefix_end = _find_section_prefix_end(
                    txt,
                    [matched_canonical or "", *section_aliases_for_report().get(matched_canonical or "", []), matched_title or ""],
                )
                if prefix_end is not None:
                    trim_paragraph_leading_text(p, prefix_end)
            else:
                _replace_paragraph_with_section_banner(
                    p,
                    img_path=img_path,
                    section_banner_width_cm=section_banner_width_cm,
                    section_tag=section_tag,
                )
            inserted += 1
        except Exception as exc:
            elog(f"Section banner failed for '{txt}': {img_path} ({exc})")
            continue

    return inserted


def _insert_question_difficulty_table(
    doc: Document,
    section_item: dict,
    column_width_cm: float,
    section_banner_map: dict[str, str] | None = None,
    section_banner_width_cm: float | None = None,
    include_answer_key: bool = False,
    add_chapter_performance: bool = False,
) -> int:
    questoes = section_item.get("questoes") or []
    if not questoes:
        return 0

    secao = section_item["secao"]
    title_p = doc.add_paragraph("")

    inserted_banner = False
    norm_secao = normalize_text_key(secao)
    matched_img = None
    if section_banner_map:
        matched_img = section_banner_map.get(norm_secao)
        if matched_img is None:
            for norm_title, img in section_banner_map.items():
                if _matches_section_title(norm_secao, norm_title):
                    matched_img = img
                    break

    if matched_img:
        img_path = resolve_path(matched_img)
        if img_path.exists():
            try:
                banner_width = float(section_banner_width_cm or column_width_cm)
                section_tag = f"{SECTION_TAG}_AUTO_{safe_tag_suffix(norm_secao)}"
                _replace_paragraph_with_section_banner(
                    title_p,
                    img_path=img_path,
                    section_banner_width_cm=banner_width,
                    section_tag=section_tag,
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                )
                inserted_banner = True
            except Exception as exc:
                elog(f"Autoavaliação banner failed for '{secao}': {img_path} ({exc})")
        else:
            elog(f"Autoavaliação banner missing for '{secao}': {img_path}")

    if not inserted_banner:
        title_p.add_run(f"Autoavaliação - {secao}")
        title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if title_p.runs:
            title_p.runs[0].bold = True
            title_p.runs[0].font.size = Pt(11)
            title_p.runs[0].font.color.rgb = RGBColor(31, 41, 55)

    fmt = title_p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if include_answer_key:
        table = doc.add_table(rows=1 + len(questoes), cols=6)
    else:
        table = doc.add_table(rows=2 + len(questoes), cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = False
    table_width_cm = max(5.0, float(column_width_cm))
    set_table_width(table, table_width_cm)

    # Proporções:
    # Com gabarito: Questão | Nível | Gabarito | Acertei | Errei | Revisar
    # Sem gabarito: Questão | Dificuldade | Acertei | Errou | Revisar
    if include_answer_key:
        # Coluna "Nível" fixa em 1,8 cm; demais colunas dividem o espaço restante.
        nivel_width_cm = 1.8
        other_base_ratios = [0.24, 0.16, 0.14, 0.14, 0.14]  # Q, G, A, E, R
        available_cm = max(1.0, table_width_cm - nivel_width_cm)
        base_sum = sum(other_base_ratios) or 1.0
        q_w = available_cm * (other_base_ratios[0] / base_sum)
        g_w = available_cm * (other_base_ratios[1] / base_sum)
        a_w = available_cm * (other_base_ratios[2] / base_sum)
        e_w = available_cm * (other_base_ratios[3] / base_sum)
        r_w = available_cm * (other_base_ratios[4] / base_sum)
        col_widths = [q_w, nivel_width_cm, g_w, a_w, e_w, r_w]
    else:
        col_ratios = [0.30, 0.24, 0.15, 0.15, 0.16]
        col_widths = [table_width_cm * ratio for ratio in col_ratios]
    for row in table.rows:
        for idx, width_cm in enumerate(col_widths):
            set_cell_width(row.cells[idx], width_cm)

    if include_answer_key:
        headers = ["Questão", "Nível", "Gabarito", "Acertei", "Errei", "Revisar"]
    else:
        headers = ["Questão", "Dificuldade", "Acertou", "Errou", "Revisar"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_fill(cell, "1F2937")
        header_size = 10
        header_bold = True
        if include_answer_key and col in (2, 3, 4, 5):
            # Evita quebra de "Gabarito/Acertei/Errei/Revisar".
            set_cell_no_wrap(cell, True)
            header_bold = False
        style_cell_text(
            cell,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=header_bold,
            color_hex="FFFFFF",
            size_pt=header_size,
        )

    counts = {"facil": 0, "media": 0, "dificil": 0, "neutro": 0}
    for i, q in enumerate(questoes, start=1):
        diff = q.get("dificuldade") or "neutro"
        counts[diff] = counts.get(diff, 0) + 1
        row = table.rows[i]
        row.cells[0].text = q["questao"]
        row.cells[1].text = "" if diff == "neutro" else difficulty_label(diff)

        if include_answer_key:
            row.cells[2].text = str(q.get("gabarito") or "")
            row.cells[3].text = "☐"
            row.cells[4].text = "☐"
            row.cells[5].text = "☐"
        else:
            row.cells[2].text = "☐"
            row.cells[3].text = "☐"
            row.cells[4].text = "☐"

        # Zebra rows para leitura.
        if i % 2 == 0:
            for c in row.cells:
                set_cell_fill(c, "F8FAFC")

        diff_fill = {
            "facil": "DCFCE7",
            "media": "FEF3C7",
            "dificil": "FEE2E2",
            "neutro": "E5E7EB",
        }.get(diff, "EEF2FF")
        set_cell_fill(row.cells[1], diff_fill)

        style_cell_text(row.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, size_pt=10)
        style_cell_text(row.cells[1], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, size_pt=10)
        if include_answer_key:
            style_cell_text(
                row.cells[2],
                align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                bold=False,
                size_pt=10,
            )
            for col in (3, 4, 5):
                style_cell_text(
                    row.cells[col],
                    align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    bold=False,
                    size_pt=10,
                )
        else:
            for col in (2, 3, 4):
                style_cell_text(
                    row.cells[col],
                    align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    bold=True,
                    color_hex="374151",
                    size_pt=12,
                )

    if not include_answer_key:
        # Linha de resumo visual no rodapé da tabela.
        summary_idx = len(questoes) + 1
        srow = table.rows[summary_idx]
        srow.cells[0].text = "Resumo"
        srow.cells[1].text = f"Fácil {counts['facil']} | Média {counts['media']} | Difícil {counts['dificil']}"
        srow.cells[2].text = "Acertos: ____"
        srow.cells[3].text = "Erros: ____"
        srow.cells[4].text = "Revisar: ____"
        for c in srow.cells:
            set_cell_fill(c, "E5E7EB")
        style_cell_text(srow.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
        style_cell_text(srow.cells[1], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
        style_cell_text(srow.cells[2], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
        style_cell_text(srow.cells[3], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
        style_cell_text(srow.cells[4], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

    if add_chapter_performance:
        summary_cols = 6 if include_answer_key else 5
        summary = doc.add_table(rows=1, cols=summary_cols)
        try:
            summary.style = "Table Grid"
        except Exception:
            pass
        summary.autofit = False
        set_table_width(summary, table_width_cm)

        summary_widths = col_widths
        row = summary.rows[0]
        for idx, width_cm in enumerate(summary_widths):
            set_cell_width(row.cells[idx], width_cm)

        row.cells[0].text = "Resumo"
        if include_answer_key:
            row.cells[1].text = ""
            row.cells[2].text = ""
            row.cells[3].text = "Acertos: ___"
            row.cells[4].text = "Erros: ___"
            row.cells[5].text = "Revisar: ___"
            for col in (2, 3, 4, 5):
                set_cell_no_wrap(row.cells[col], True)
        else:
            row.cells[1].text = ""
            row.cells[2].text = "Acertos: ___"
            row.cells[3].text = "Erros: ___"
            row.cells[4].text = "Revisar: ___"

        for c in row.cells:
            set_cell_fill(c, "E5E7EB")
        style_cell_text(row.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True, size_pt=10)
        for col in range(1, summary_cols):
            is_plain_col = include_answer_key and col in (2, 3, 4, 5)
            style_cell_text(
                row.cells[col],
                align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                bold=not is_plain_col,
                size_pt=10,
            )

    # Espaço entre tabelas no bloco final de autoavaliação.
    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)

    return 1


def insert_question_difficulty_tables(
    doc: Document,
    sections_data: list[dict],
    column_width_cm: float,
    section_banners: dict[str, str] | None = None,
    section_banner_width_cm: float | None = None,
    include_answer_key: bool = False,
    add_chapter_performance: bool = False,
    single_column_section: bool = True,
) -> int:
    inserted = 0
    if not sections_data:
        return inserted

    effective_width_cm = float(column_width_cm)
    if single_column_section:
        try:
            sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
            _set_section_single_column(sec)
            text_width_cm = (
                float(sec.page_width.cm)
                - float(sec.left_margin.cm)
                - float(sec.right_margin.cm)
            )
            # Usa largura real da página para evitar tabela espremida em layout de 2 colunas.
            effective_width_cm = max(effective_width_cm, max(8.0, text_width_cm))
        except Exception as exc:
            elog(f"Failed to create single-column section for autoavaliação: {exc}")

    section_banner_map: dict[str, str] = {}
    if isinstance(section_banners, dict):
        for title, img in section_banners.items():
            norm_title = normalize_text_key(title)
            if norm_title:
                section_banner_map[norm_title] = img

    # Todas as autoavaliações ficam no final do documento, na ordem das seções.
    block_title = doc.add_paragraph("Quadro de Autoavaliação por Seção")
    block_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if block_title.runs:
        block_title.runs[0].bold = True
        block_title.runs[0].font.size = Pt(12)
        block_title.runs[0].font.color.rgb = RGBColor(17, 24, 39)
    bfmt = block_title.paragraph_format
    bfmt.space_before = Pt(10)
    bfmt.space_after = Pt(8)
    bfmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for item in sections_data:
        inserted += _insert_question_difficulty_table(
            doc,
            item,
            column_width_cm=effective_width_cm,
            section_banner_map=section_banner_map,
            section_banner_width_cm=section_banner_width_cm,
            include_answer_key=include_answer_key,
            add_chapter_performance=add_chapter_performance,
        )
    return inserted


def append_difficulty_report_appendix(
    doc: Document,
    report: dict,
    *,
    margin_top_cm: float = 2.0,
    margin_bottom_cm: float = 2.0,
    margin_left_cm: float = 2.5,
    margin_right_cm: float = 2.5,
) -> bool:
    sections = report.get("secoes") if isinstance(report, dict) else None
    if not sections:
        return False

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    try:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(margin_left_cm)
        section.right_margin = Cm(margin_right_cm)
        section.top_margin = Cm(margin_top_cm)
        section.bottom_margin = Cm(margin_bottom_cm)
    except Exception:
        pass

    try:
        text_width_cm = (
            float(section.page_width.cm)
            - float(section.left_margin.cm)
            - float(section.right_margin.cm)
        )
    except Exception:
        text_width_cm = 17.4
    text_width_cm = max(10.0, text_width_cm)

    title = doc.add_paragraph("Relatório de Dificuldade por Seção")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if title.runs:
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(13)
        title.runs[0].font.color.rgb = RGBColor(17, 24, 39)
    tfmt = title.paragraph_format
    tfmt.space_before = Pt(0)
    tfmt.space_after = Pt(6)

    meta = [
        f"Conteúdo: {report.get('conteudo', '-')}",
        f"Área: {report.get('area_conhecimento', '-')}",
    ]
    for line in meta:
        p = doc.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(75, 85, 99)

    totals = report.get("totais") or {"facil": 0, "media": 0, "dificil": 0, "total": 0}
    cards = doc.add_table(rows=2, cols=4)
    cards.autofit = False
    set_table_width(cards, text_width_cm)
    card_width_cm = text_width_cm / 4
    for row in cards.rows:
        for c in row.cells:
            set_cell_width(c, card_width_cm)
    card_headers = ["Fácil", "Média", "Difícil", "Total"]
    card_values = [totals["facil"], totals["media"], totals["dificil"], totals["total"]]
    header_fill = ["DCFCE7", "FEF3C7", "FEE2E2", "E5E7EB"]
    value_color = ["166534", "92400E", "991B1B", "111827"]
    for idx in range(4):
        cards.cell(0, idx).text = card_headers[idx]
        set_cell_fill(cards.cell(0, idx), header_fill[idx])
        style_cell_text(
            cards.cell(0, idx),
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            size_pt=9,
        )
        cards.cell(1, idx).text = str(card_values[idx])
        style_cell_text(
            cards.cell(1, idx),
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            color_hex=value_color[idx],
            size_pt=12,
        )

    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=2 + len(sections), cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = False
    set_table_width(table, text_width_cm)

    col_ratios = [0.52, 0.12, 0.12, 0.12, 0.12]
    col_widths = [text_width_cm * ratio for ratio in col_ratios]
    for row in table.rows:
        for idx, width_cm in enumerate(col_widths):
            set_cell_width(row.cells[idx], width_cm)

    headers = ["Seção", "Fácil", "Média", "Difícil", "Total"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_fill(cell, "1F2937")
        style_cell_text(
            cell,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            color_hex="FFFFFF",
            size_pt=9,
        )

    for idx, row_data in enumerate(sections, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(row_data["secao"])
        row.cells[1].text = str(row_data["facil"])
        row.cells[2].text = str(row_data["media"])
        row.cells[3].text = str(row_data["dificil"])
        row.cells[4].text = str(row_data["total"])

        if idx % 2 == 0:
            for c in row.cells:
                set_cell_fill(c, "F8FAFC")

        style_cell_text(row.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True, size_pt=9)
        style_cell_text(row.cells[1], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="166534", size_pt=9)
        style_cell_text(row.cells[2], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="92400E", size_pt=9)
        style_cell_text(row.cells[3], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="991B1B", size_pt=9)
        style_cell_text(row.cells[4], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="111827", size_pt=9)

    total_row_idx = len(sections) + 1
    total_row = table.rows[total_row_idx]
    total_row.cells[0].text = "TOTAL"
    total_row.cells[1].text = str(totals["facil"])
    total_row.cells[2].text = str(totals["media"])
    total_row.cells[3].text = str(totals["dificil"])
    total_row.cells[4].text = str(totals["total"])
    for c in total_row.cells:
        set_cell_fill(c, "E5E7EB")
        style_cell_text(
            c,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            color_hex="111827",
            size_pt=9,
        )
    style_cell_text(
        total_row.cells[0],
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        bold=True,
        color_hex="111827",
        size_pt=9,
    )

    return True
