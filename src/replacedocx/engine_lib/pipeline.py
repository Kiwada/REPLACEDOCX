from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

from .common import (
    ALTERNATIVE_ONLY_RE,
    BADGE_TAG,
    GABARITO_RE,
    default_markers_for_area,
    default_section_banners_for_area,
    elog,
    normalize_text_key,
)
from .docx_utils import (
    apply_paragraph_layout,
    apply_run_font,
    is_figure_caption_text,
    is_reference_caption_text,
    is_reference_or_citation_text,
    iter_paragraphs,
    paragraph_has_drawing,
    remove_paragraph,
    replace_question_prefix_marker_in_paragraph,
)
from .renderers import (
    append_difficulty_report_appendix,
    apply_section_banners,
    insert_question_difficulty_tables,
)
from .report import (
    collect_questions_by_section,
    gerar_relatorio_dificuldade_por_secao,
    match_section_for_report,
)

ANSWER_KEY_ONLY_LINE_RE = re.compile(
    r"^\s*\d{1,3}\s*[:\)\.\-]\s*(?:ALTERNATIVA\s*)?(?:\[|\()?\s*[A-E]\s*(?:\]|\))?\s*$",
    re.IGNORECASE,
)


def _is_reference_below_image(paragraphs: list, idx: int, lookback: int = 3) -> bool:
    if idx < 0 or idx >= len(paragraphs):
        return False
    current = paragraphs[idx]
    if paragraph_has_drawing(current):
        return True

    lower = max(-1, idx - lookback - 1)
    for j in range(idx - 1, lower, -1):
        prev = paragraphs[j]
        if paragraph_has_drawing(prev):
            return True
        prev_text = (prev.text or "").strip()
        if prev_text and not (
            is_figure_caption_text(prev_text) or is_reference_caption_text(prev_text)
        ):
            return False
    return False


def _apply_reference_caption_font(
    doc: Document,
    reference_font_name: str = "Arial",
    reference_font_size: int = 8,
) -> int:
    updated = 0
    paragraph_groups = [list(doc.paragraphs)]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraph_groups.append(list(cell.paragraphs))

    for paragraphs in paragraph_groups:
        for i, p in enumerate(paragraphs):
            raw_text = (p.text or "").strip()
            if not is_reference_or_citation_text(raw_text):
                continue

            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if not p.runs:
                p.add_run("")
            for run in p.runs:
                apply_run_font(run, reference_font_name, reference_font_size)
            updated += 1

    return updated


def _replace_question_markers_in_exercise_sections(
    doc: Document,
    markers: dict[str, str],
    badge_width_cm: float,
    font_name: str,
    font_size: int,
) -> int:
    replaced = 0
    in_exercise_block = False

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            section = match_section_for_report(normalize_text_key(txt))
            if section:
                in_exercise_block = True
                continue

        if not in_exercise_block:
            continue

        if replace_question_prefix_marker_in_paragraph(
            p,
            markers,
            badge_width_cm,
            font_name,
            font_size,
            badge_tag=BADGE_TAG,
        ):
            replaced += 1

    return replaced


def _remove_tail_answer_key_appendix(doc: Document) -> int:
    paragraphs = list(doc.paragraphs)
    if len(paragraphs) < 20:
        return 0

    tail_start = int(len(paragraphs) * 0.5)
    start_idx: int | None = None

    for i in range(tail_start, len(paragraphs)):
        txt = (paragraphs[i].text or "").strip()
        if not txt:
            continue

        norm_txt = normalize_text_key(txt)
        if not match_section_for_report(norm_txt):
            continue

        answers_found = 0
        scanned = 0
        j = i + 1
        while j < len(paragraphs) and scanned < 60:
            next_txt = (paragraphs[j].text or "").strip()
            if not next_txt:
                j += 1
                continue

            scanned += 1
            next_norm = normalize_text_key(next_txt)
            if ANSWER_KEY_ONLY_LINE_RE.match(next_norm):
                answers_found += 1
            elif match_section_for_report(next_norm):
                pass
            elif next_norm.startswith(("CAPITULO", "AULA", "GABARITO", "RESPOSTA")):
                pass
            elif answers_found > 0:
                break

            j += 1

        if answers_found >= 3:
            start_idx = i
            for k in range(max(tail_start, i - 3), i):
                prev_norm = normalize_text_key((paragraphs[k].text or "").strip())
                if prev_norm.startswith(("CAPITULO", "AULA", "GABARITO", "RESPOSTA")):
                    start_idx = k
                    break
            break

    if start_idx is None:
        return 0

    removed = 0
    for idx in range(len(paragraphs) - 1, start_idx - 1, -1):
        remove_paragraph(paragraphs[idx])
        removed += 1
    return removed


def processar_docx(input_path: str | Path, output_path: str | Path, config: dict):
    elog("\n============================")
    elog("START " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {input_path}")

    font_name = (config.get("font_name") or "Arial").strip()
    font_size = int(config.get("font_size", 11))
    badge_width_cm = float(config.get("badge_width_cm", 1.3))
    column_width_cm = float(config.get("column_width_cm", 7.5))
    remove_gabarito = bool(config.get("remove_gabarito", True))
    justify = bool(config.get("justify", True))
    format_text = bool(config.get("format_text", True))
    area_conhecimento = (config.get("area_conhecimento") or "biologia").strip()
    insert_section_banners = bool(config.get("insert_section_banners", True))
    insert_question_tables = bool(config.get("insert_question_tables", True))
    add_section_summary_row = bool(config.get("add_section_summary_row", True))
    append_difficulty_report = bool(config.get("append_difficulty_report", False))
    difficulty_report_data = config.get("difficulty_report_data")
    if difficulty_report_data is not None and not isinstance(difficulty_report_data, dict):
        raise ValueError("Config inválida: 'difficulty_report_data' deve ser um dict.")
    margin_top_cm = float(config.get("margin_top_cm", 2.0))
    margin_bottom_cm = float(config.get("margin_bottom_cm", 2.0))
    margin_left_cm = float(config.get("margin_left_cm", 2.5))
    margin_right_cm = float(config.get("margin_right_cm", 2.5))
    section_banner_width_raw = config.get("section_banner_width_cm")
    section_banner_width_cm = float(
        section_banner_width_raw if section_banner_width_raw is not None else column_width_cm
    )

    base_markers = default_markers_for_area(area_conhecimento)
    user_markers = config.get("markers") or {}
    if user_markers and not isinstance(user_markers, dict):
        raise ValueError("Config inválida: 'markers' deve ser um dict.")
    markers = dict(base_markers)
    markers.update(user_markers)
    if not markers:
        raise ValueError("Config inválida: 'markers' deve ser um dict com pelo menos 1 item.")

    base_sections = default_section_banners_for_area(area_conhecimento)
    user_sections = config.get("section_banners") or {}
    if user_sections and not isinstance(user_sections, dict):
        raise ValueError("Config inválida: 'section_banners' deve ser um dict.")
    section_banners = dict(base_sections)
    section_banners.update(user_sections)

    finalize_word = bool(config.get("finalize_word", True))
    force_inline_wrap = bool(config.get("force_inline_wrap", True))

    doc = Document(str(input_path))
    for section in doc.sections:
        section.top_margin = Cm(margin_top_cm)
        section.bottom_margin = Cm(margin_bottom_cm)
        section.left_margin = Cm(margin_left_cm)
        section.right_margin = Cm(margin_right_cm)

    if format_text:
        for p in iter_paragraphs(doc):
            apply_paragraph_layout(p, justify)
            for r in p.runs:
                apply_run_font(r, font_name, font_size)
    else:
        elog("Skipped global text formatting (format_text=False).")

    sections_data = (
        collect_questions_by_section(doc, include_answer_key=True)
        if insert_question_tables
        else []
    )
    if sections_data:
        total_q = sum(len(item.get("questoes") or []) for item in sections_data)
        total_g = sum(
            1
            for item in sections_data
            for q in (item.get("questoes") or [])
            if q.get("gabarito")
        )
        elog(f"Answer key mapped: {total_g}/{total_q}")
        removed_tail = _remove_tail_answer_key_appendix(doc)
        if removed_tail:
            elog(f"Removed tail answer-key appendix paragraphs: {removed_tail}")

    if remove_gabarito:
        paragraphs = list(iter_paragraphs(doc))
        to_delete_idx: set[int] = set()

        for i, p in enumerate(paragraphs):
            txt = (p.text or "").strip()
            m = GABARITO_RE.match(txt)
            if not m:
                continue
            to_delete_idx.add(i)

            # Caso comum:
            # "Resposta:" em uma linha e alternativa ("A", "B", ...) na próxima.
            if m.groupdict().get("alt"):
                continue

            j = i + 1
            while j < len(paragraphs):
                next_txt = (paragraphs[j].text or "").strip()
                if not next_txt:
                    j += 1
                    continue
                if ALTERNATIVE_ONLY_RE.match(next_txt):
                    to_delete_idx.add(j)
                break

        for idx in sorted(to_delete_idx, reverse=True):
            remove_paragraph(paragraphs[idx])

        elog("Removed answer lines (gabarito/resposta): " + str(len(to_delete_idx)))

    replaced_badges = _replace_question_markers_in_exercise_sections(
        doc,
        markers,
        badge_width_cm,
        font_name,
        font_size,
    )
    elog("Replaced difficulty markers in exercise sections: " + str(replaced_badges))

    if insert_section_banners and section_banners:
        inserted = apply_section_banners(
            doc,
            section_banners=section_banners,
            section_banner_width_cm=section_banner_width_cm,
        )
        elog("Inserted section banners: " + str(inserted))

    ref_count = _apply_reference_caption_font(
        doc,
        reference_font_name="Arial",
        reference_font_size=8,
    )
    if ref_count:
        elog("Applied Arial 8 and right alignment to references/citations: " + str(ref_count))

    if insert_question_tables and sections_data:
        inserted_tables = insert_question_difficulty_tables(
            doc,
            sections_data,
            column_width_cm=column_width_cm,
            section_banners=section_banners,
            section_banner_width_cm=section_banner_width_cm,
            include_answer_key=True,
            add_chapter_performance=add_section_summary_row,
            single_column_section=True,
        )
        elog("Inserted question difficulty tables: " + str(inserted_tables))

    if append_difficulty_report:
        report_data = difficulty_report_data or gerar_relatorio_dificuldade_por_secao(
            input_path,
            area_conhecimento=area_conhecimento,
        )
        if append_difficulty_report_appendix(
            doc,
            report_data,
            margin_top_cm=margin_top_cm,
            margin_bottom_cm=margin_bottom_cm,
            margin_left_cm=margin_left_cm,
            margin_right_cm=margin_right_cm,
        ):
            elog("Inserted A4 difficulty report appendix at document end.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    elog("Saved via python-docx OK")

    if finalize_word:
        from ..word_finalize import finalize_with_word

        finalize_with_word(
            docx_in=output_path,
            docx_out=output_path,
            column_width_cm=column_width_cm,
            font_name=font_name,
            font_size=font_size,
            justify=justify,
            force_inline_wrap=force_inline_wrap,
            badge_tag=BADGE_TAG,
        )
        elog("Finalized with Word OK")

    return {"arquivo_saida": str(output_path)}
