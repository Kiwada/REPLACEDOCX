from __future__ import annotations

import hashlib
import io
import os
import re
import subprocess
import sys
import unicodedata
from pathlib import Path

from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Cm

GABARITO_RE = re.compile(
    r"^\s*(?:GABARITO|RESPOSTA(?:\s+CORRETA)?)\s*[:\-]?\s*(?:ALTERNATIVA\s*)?(?P<alt>[A-E])?\s*[\)\.\-]?\s*$",
    re.IGNORECASE,
)
ALTERNATIVE_ONLY_RE = re.compile(r"^\s*(?:ALTERNATIVA\s*)?(?P<alt>[A-E])\s*[\)\.\-]?\s*$", re.IGNORECASE)
QUESTION_DIFFICULTY_RE = re.compile(
    r"^\s*(?:(?P<num>\d+)\s*[\.\)\-:]?\s*)?(?:NIVEL\s*[:\-]?\s*)?\(?\s*(?P<level>FACIL|MEDIA|MEDIO|MEDIAS|DIFICIL)\s*\)?\s*[:\-\.]?\s*$"
)
BADGE_TAG = "BADGE_REPLACE_DOCX"
SECTION_TAG = "SECTION_BANNER_REPLACE_DOCX"


def runtime_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    # common.py está em .../src/replacedocx/engine_lib/common.py
    # runtime da engine deve seguir na raiz do projeto.
    return Path(__file__).resolve().parents[3]


ENGINE_LOG = runtime_dir() / "engine_debug.log"


def elog(msg: str) -> None:
    try:
        with open(ENGINE_LOG, "a", encoding="utf-8") as f:
            f.write(msg.rstrip() + "\n")
    except Exception:
        return


def assets_dir() -> Path:
    # Portabilidade: mantém Windows e adiciona caminho padrão para macOS/Linux.
    if sys.platform == "darwin":
        base = Path.home() / "Library" / "Application Support"
    elif os.name == "nt":
        base = Path(os.environ.get("APPDATA") or (Path.home() / "AppData" / "Roaming"))
    else:
        base = Path(os.environ.get("XDG_DATA_HOME") or (Path.home() / ".local" / "share"))
    return base / "ReplaceDocx" / "assets"


def normalize_area_slug(area: str) -> str:
    txt = unicodedata.normalize("NFD", (area or "").strip().lower())
    txt = "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")
    txt = re.sub(r"[^a-z0-9]+", "_", txt).strip("_")
    return txt or "geral"


def normalize_text_key(text: str) -> str:
    txt = unicodedata.normalize("NFD", (text or "").strip().upper())
    txt = "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")
    txt = re.sub(r"\s+", " ", txt)
    return txt


def _remove_accents(text: str) -> str:
    txt = unicodedata.normalize("NFD", text or "")
    return "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")


def _difficulty_variants(label: str) -> list[str]:
    base = label.strip()
    no_acc = _remove_accents(base)
    candidates = {
        base,
        base.title(),
        base.lower(),
        no_acc,
        no_acc.title(),
        no_acc.lower(),
    }
    ordered = sorted((c for c in candidates if c), key=lambda s: (len(s), s))
    return ordered


def safe_tag_suffix(text: str) -> str:
    key = normalize_text_key(text)
    return re.sub(r"[^A-Z0-9]+", "_", key).strip("_") or "GENERIC"


def default_markers_for_area(area: str) -> dict[str, str]:
    area_slug = normalize_area_slug(area)
    base = f"areas/{area_slug}/capsulas"
    levels = [
        (f"{base}/facil.png", ["FÁCIL"]),
        (f"{base}/media.png", ["MÉDIA", "MÉDIO", "MÉDIAS"]),
        (f"{base}/dificil.png", ["DIFÍCIL"]),
    ]
    markers: dict[str, str] = {}
    for img, labels in levels:
        for label in labels:
            for variant in _difficulty_variants(label):
                markers[variant] = img
                markers[f"({variant})"] = img
    return markers


def default_section_banners_for_area(area: str) -> dict[str, str]:
    area_slug = normalize_area_slug(area)
    base = f"areas/{area_slug}/secoes"
    jpg_base = f"{base}/JPG"

    # Algumas áreas organizam seções em subpasta JPG.
    project_jpg_dir = runtime_dir() / "Assets" / "areas" / area_slug / "secoes" / "JPG"
    user_jpg_dir = assets_dir() / "areas" / area_slug / "secoes" / "JPG"
    if project_jpg_dir.exists() or user_jpg_dir.exists():
        base = jpg_base
    banners = {
        "EXERCÍCIOS DE SALA": f"{base}/exercicios_sala.png",
        "QUESTÕES DE SALA": f"{base}/exercicios_sala.png",
        "QUESTÃO DE SALA": f"{base}/exercicios_sala.png",
        "EXERCÍCIOS PROPOSTOS": f"{base}/exercicios_propostos.png",
        "QUESTÕES PROPOSTAS": f"{base}/exercicios_propostos.png",
        "QUESTÃO PROPOSTA": f"{base}/exercicios_propostos.png",
        "SEÇÃO ENEM": f"{base}/secao_enem.png",
        "QUESTÕES ENEM": f"{base}/secao_enem.png",
        "QUESTÃO ENEM": f"{base}/secao_enem.png",
        "EXERCÍCIOS DE APROFUNDAMENTO": f"{base}/exercicios_aprofundamento.png",
        "QUESTÕES DE APROFUNDAMENTO": f"{base}/exercicios_aprofundamento.png",
        "QUESTÃO DE APROFUNDAMENTO": f"{base}/exercicios_aprofundamento.png",
        "EXERCÍCIOS REGIONAIS": f"{base}/exercicios_regionais.png",
        "QUESTÕES REGIONAIS": f"{base}/exercicios_regionais.png",
        "QUESTÃO REGIONAL": f"{base}/exercicios_regionais.png",
        "EXERCÍCIO REGIONAL": f"{base}/exercicios_regionais.png",
    }

    # Em matemática, essa seção não é usada no material atual.
    if area_slug != "matematica":
        banners.update(
            {
                "EXERCÍCIO DISSERTATIVO": f"{base}/exercicios_dissertativos.png",
                "EXERCÍCIOS DISSERTATIVOS": f"{base}/exercicios_dissertativos.png",
            }
        )

    return banners


def paths_with_image_extension_fallback(p: Path) -> list[Path]:
    # Prioridade: PNG primeiro para preservar cor/transparência de assets finais.
    preferred_exts = [".png", ".jpg", ".jpeg"]
    image_exts = {".jpg", ".jpeg", ".png"}

    if p.suffix.lower() in image_exts:
        stem = p.with_suffix("")
        return [stem.with_suffix(ext) for ext in preferred_exts]

    if p.suffix:
        return [p]

    return [p.with_suffix(ext) for ext in preferred_exts]


def convert_image_for_docx(img_path: Path) -> Path:
    if img_path.suffix.lower() not in {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}:
        return img_path

    cache_dir = runtime_dir() / ".image_cache_docx"
    cache_dir.mkdir(parents=True, exist_ok=True)
    # Inclui versão do pipeline para invalidar cache antigo quando
    # a lógica de conversão/colorimetria for aprimorada.
    key = f"v4_srgb_embed_any::{img_path.resolve()}::{img_path.stat().st_mtime_ns}::{img_path.stat().st_size}"
    out_name = hashlib.sha1(key.encode("utf-8")).hexdigest() + ".png"
    out_path = cache_dir / out_name
    if out_path.exists():
        return out_path

    # Tenta Pillow primeiro (se existir no ambiente), com conversão para sRGB.
    try:
        from PIL import Image, ImageCms  # type: ignore

        with Image.open(str(img_path)) as im:
            src_icc = im.info.get("icc_profile")
            srgb_icc = None
            try:
                srgb_profile = ImageCms.createProfile("sRGB")
                srgb_icc = ImageCms.ImageCmsProfile(srgb_profile).tobytes()
            except Exception:
                srgb_profile = None

            if src_icc:
                try:
                    src_profile = ImageCms.ImageCmsProfile(io.BytesIO(src_icc))
                    if srgb_profile is not None:
                        im = ImageCms.profileToProfile(
                            im,
                            src_profile,
                            srgb_profile,
                            outputMode="RGBA" if im.mode == "RGBA" else "RGB",
                        )
                except Exception:
                    pass

            if im.mode not in {"RGB", "RGBA"}:
                im = im.convert("RGB")
            save_kwargs = {"format": "PNG", "optimize": True}
            if srgb_icc:
                save_kwargs["icc_profile"] = srgb_icc
            im.save(str(out_path), **save_kwargs)
        if out_path.exists():
            return out_path
    except Exception:
        pass

    # Fallback macOS via sips.
    if sys.platform == "darwin":
        try:
            subprocess.run(
                ["sips", "-s", "format", "png", str(img_path), "--out", str(out_path)],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            if out_path.exists():
                return out_path
        except Exception:
            pass

    return img_path


def add_picture_resilient(run, img_path: Path, width_cm: float):
    if img_path.suffix.lower() in {".jpg", ".jpeg"}:
        converted = convert_image_for_docx(img_path)
        if converted.exists():
            img_path = converted

    try:
        return run.add_picture(str(img_path), width=Cm(width_cm))
    except UnrecognizedImageError:
        converted = convert_image_for_docx(img_path)
        if converted != img_path and converted.exists():
            return run.add_picture(str(converted), width=Cm(width_cm))
        raise


def resolve_path(p: str | Path) -> Path:
    p = Path(p)

    if p.is_absolute() and p.exists():
        return p

    assets = assets_dir()
    base = runtime_dir()
    project_assets = base / "Assets"
    candidates = [
        (assets / p),
        (assets / p.name),
        (base / p),
        (base / "Assets" / p),
        (base / "Assets" / p.name),
        (base / "assets" / p),
        (base / "assets" / p.name),
        (project_assets / p),
        (project_assets / p.name),
    ]
    expanded: list[Path] = []
    for cand in candidates:
        expanded.extend(paths_with_image_extension_fallback(cand))

    # Fallback: alguns projetos organizam assets em subpastas de 1 nível
    # (ex.: .../secoes/JPG/arquivo.png). Tenta localizar automaticamente.
    nested_expanded: list[Path] = []
    for cand in list(expanded):
        parent = cand.parent
        if not parent.exists():
            continue
        try:
            for child in parent.iterdir():
                if not child.is_dir():
                    continue
                nested_expanded.extend(paths_with_image_extension_fallback(child / cand.name))
        except Exception:
            continue
    expanded.extend(nested_expanded)

    seen = set()
    for cand in expanded:
        cand_resolved = cand.resolve()
        key = str(cand_resolved)
        if key in seen:
            continue
        seen.add(key)
        if cand_resolved.exists():
            return cand_resolved

    return (base / p).resolve()
