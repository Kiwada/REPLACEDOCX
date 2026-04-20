from __future__ import annotations

import re
import unicodedata

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .common import BADGE_TAG, add_picture_resilient, resolve_path

QUESTION_PREFIX_MARKER_RE = re.compile(
    r"^\s*(?:(?P<num_a>\d+)\s*[\.\)\-:]?\s*(?P<marker_a>\(\s*[^()]{2,32}\s*\)|[^\s()]{2,20})|(?P<marker_b>\(\s*[^()]{2,32}\s*\)|[^\s()]{2,20})\s*(?P<num_b>\d+)\s*[\.\)\-:]?)"
)
PAREN_CITATION_RE = re.compile(r"^\(\s*.+\b(?:18|19|20)\d{2}\s*\)\.?$")
INLINE_CITATION_RE = re.compile(r"^SEGUNDO\s+.+\((?:18|19|20)\d{2}\)")
BIBLIOGRAPHIC_ENTRY_RE = re.compile(
    r"^[A-Z][A-Z'`\-]+(?:\s+[A-Z][A-Z'`\-]+)*(?:\s+ET\s+AL\.?)?(?:\s*;\s*[A-Z][A-Z'`\-]+(?:\s+[A-Z][A-Z'`\-]+)*)*,\s+.+\b(?:18|19|20)\d{2}\.?$"
)


def set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def set_cell_no_wrap(cell, no_wrap: bool = True) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap_el = tc_pr.find(qn("w:noWrap"))
    if no_wrap:
        if no_wrap_el is None:
            no_wrap_el = OxmlElement("w:noWrap")
            tc_pr.append(no_wrap_el)
        no_wrap_el.set(qn("w:val"), "1")
        return

    if no_wrap_el is not None:
        tc_pr.remove(no_wrap_el)


def set_cell_fit_text(cell, fit_text: bool = True) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    fit_text_el = tc_pr.find(qn("w:tcFitText"))
    if fit_text:
        if fit_text_el is None:
            fit_text_el = OxmlElement("w:tcFitText")
            tc_pr.append(fit_text_el)
        fit_text_el.set(qn("w:val"), "1")
        return

    if fit_text_el is not None:
        tc_pr.remove(fit_text_el)


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def style_cell_text(
    cell,
    *,
    align=WD_PARAGRAPH_ALIGNMENT.LEFT,
    bold: bool = False,
    color_hex: str | None = None,
    size_pt: int = 10,
    font_name: str = "Arial",
) -> None:
    if not cell.paragraphs:
        return
    p = cell.paragraphs[0]
    p.alignment = align
    if not p.runs:
        p.add_run("")
    for run in p.runs:
        run.bold = bold
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:cs"), font_name)
        run.font.size = Pt(size_pt)
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(color_hex)


def _iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Parent type not supported: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def iter_paragraphs(doc: Document | _Cell):
    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            yield item
            continue

        for row in item.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def remove_paragraph(p) -> None:
    p._element.getparent().remove(p._element)


def apply_paragraph_layout(p, justify: bool) -> None:
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY if justify else WD_PARAGRAPH_ALIGNMENT.LEFT


def apply_run_font(run, font_name: str, font_size: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)


def is_reference_caption_text(text: str) -> bool:
    txt = (text or "").strip()
    if not txt:
        return False
    norm = unicodedata.normalize("NFD", txt.upper())
    norm = "".join(ch for ch in norm if unicodedata.category(ch) != "Mn")
    norm = re.sub(r"\s+", " ", norm)
    return bool(re.match(r"^(FONTE|REFERENCIA|REFERENCIAS|DISPONIVEL\s+EM)\b", norm))


def is_reference_or_citation_text(text: str) -> bool:
    txt = (text or "").strip()
    if not txt:
        return False

    if is_reference_caption_text(txt):
        return True

    norm = unicodedata.normalize("NFD", txt.upper())
    norm = "".join(ch for ch in norm if unicodedata.category(ch) != "Mn")
    norm = re.sub(r"\s+", " ", norm)

    if "DISPONIVEL EM" in norm or "ACESSO EM" in norm:
        return True
    if " IN: " in norm and re.search(r"\b(?:18|19|20)\d{2}\b", norm):
        return True
    if PAREN_CITATION_RE.match(norm):
        return True
    if INLINE_CITATION_RE.match(norm):
        return True
    if BIBLIOGRAPHIC_ENTRY_RE.match(norm):
        return True
    return False


def is_figure_caption_text(text: str) -> bool:
    txt = (text or "").strip()
    if not txt:
        return False
    norm = unicodedata.normalize("NFD", txt.upper())
    norm = "".join(ch for ch in norm if unicodedata.category(ch) != "Mn")
    norm = re.sub(r"\s+", " ", norm)
    return bool(re.match(r"^(FIGURA|IMAGEM|ILUSTRACAO|FOTO|GRAFICO|QUADRO|TABELA)\b", norm))


def paragraph_has_drawing(paragraph) -> bool:
    try:
        return bool(paragraph._element.xpath(".//w:drawing"))
    except Exception:
        return False


def paragraph_has_office_math(paragraph) -> bool:
    try:
        return bool(
            paragraph._element.xpath(
                ".//*[local-name()='oMath' or local-name()='oMathPara']"
            )
        )
    except Exception:
        return False


def insert_badge_run(p, img_path, badge_width_cm: float, badge_tag: str = BADGE_TAG) -> None:
    run = p.add_run()
    pic = add_picture_resilient(run, img_path, badge_width_cm)

    # Marca a imagem para o backend do Word ignorar badges na etapa final.
    try:
        pic._inline.docPr.set("descr", badge_tag)
        pic._inline.docPr.set("title", badge_tag)
    except Exception:
        return


def _expand_markers(markers: dict[str, str]) -> dict[str, str]:
    markers_expanded: dict[str, str] = {}
    for raw_key, img in markers.items():
        for variant in (
            raw_key,
            unicodedata.normalize("NFC", raw_key),
            unicodedata.normalize("NFD", raw_key),
        ):
            if variant and variant not in markers_expanded:
                markers_expanded[variant] = img
    return markers_expanded


def _resolve_marker_key(raw_marker: str, markers_expanded: dict[str, str]) -> str | None:
    token = (raw_marker or "").strip()
    if not token:
        return None

    candidates: list[str] = [token]
    if token.startswith("(") and token.endswith(")"):
        inner = token[1:-1].strip()
        if inner:
            candidates.append(inner)
            candidates.append(f"({inner})")
    else:
        candidates.append(f"({token})")

    variants: list[str] = []
    for c in candidates:
        variants.extend((c, unicodedata.normalize("NFC", c), unicodedata.normalize("NFD", c)))

    seen: set[str] = set()
    for candidate in variants:
        if candidate in seen:
            continue
        seen.add(candidate)
        if candidate in markers_expanded:
            return candidate
    return None


def _element_visible_text(element) -> str:
    chunks: list[str] = []
    for node in element.iter():
        if node.tag == qn("w:t"):
            chunks.append(node.text or "")
        elif node.tag == qn("w:tab"):
            chunks.append("\t")
        elif node.tag in {qn("w:br"), qn("w:cr")}:
            chunks.append("\n")
    return "".join(chunks)


def _trim_element_leading_text(element, char_count: int) -> int:
    remaining = max(0, int(char_count))
    if remaining <= 0:
        return 0

    for node in list(element.iter()):
        if remaining <= 0:
            break

        if node.tag == qn("w:t"):
            text = node.text or ""
            if not text:
                continue
            if remaining >= len(text):
                remaining -= len(text)
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)
                continue
            node.text = text[remaining:]
            remaining = 0
            break

        if node.tag in {qn("w:tab"), qn("w:br"), qn("w:cr")}:
            remaining -= 1
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)

    return max(0, remaining)


def trim_paragraph_leading_text(p, char_count: int) -> None:
    remaining = max(0, int(char_count))
    if remaining <= 0:
        return

    paragraph_el = p._element
    for child in [child for child in list(paragraph_el) if child.tag != qn("w:pPr")]:
        if remaining <= 0:
            break

        child_text = _element_visible_text(child)
        child_len = len(child_text)
        if child_len <= 0:
            paragraph_el.remove(child)
            continue

        if remaining >= child_len:
            remaining -= child_len
            paragraph_el.remove(child)
            continue

        _trim_element_leading_text(child, remaining)
        break


def _detach_new_run_element(run):
    el = run._element
    el.getparent().remove(el)
    return el


def _build_text_run_element(p, text: str, font_name: str, font_size: int):
    if not text:
        return None
    run = p.add_run(text)
    apply_run_font(run, font_name, font_size)
    return _detach_new_run_element(run)


def _build_badge_run_element(
    p,
    img_path,
    badge_width_cm: float,
    badge_tag: str = BADGE_TAG,
):
    run = p.add_run()
    pic = add_picture_resilient(run, img_path, badge_width_cm)
    try:
        pic._inline.docPr.set("descr", badge_tag)
        pic._inline.docPr.set("title", badge_tag)
    except Exception:
        pass
    return _detach_new_run_element(run)


def _replace_prefix_marker_preserving_non_text(
    p,
    marker_start: int,
    marker_end: int,
    img_path,
    badge_width_cm: float,
    font_name: str,
    font_size: int,
    badge_tag: str = BADGE_TAG,
) -> None:
    paragraph_el = p._element
    original_children = [child for child in paragraph_el if child.tag != qn("w:pPr")]
    for child in original_children:
        paragraph_el.remove(child)

    text_pos = 0
    badge_inserted = False

    for child in original_children:
        child_text = _element_visible_text(child)
        if not child_text:
            if not badge_inserted and text_pos >= marker_end:
                paragraph_el.append(
                    _build_badge_run_element(
                        p,
                        img_path,
                        badge_width_cm,
                        badge_tag=badge_tag,
                    )
                )
                badge_inserted = True
            paragraph_el.append(child)
            continue

        child_start = text_pos
        child_end = child_start + len(child_text)

        if child_end <= marker_start:
            paragraph_el.append(child)
        elif child_start >= marker_end:
            if not badge_inserted:
                paragraph_el.append(
                    _build_badge_run_element(
                        p,
                        img_path,
                        badge_width_cm,
                        badge_tag=badge_tag,
                    )
                )
                badge_inserted = True
            paragraph_el.append(child)
        else:
            before_text = child_text[: max(0, marker_start - child_start)]
            after_text = child_text[max(0, marker_end - child_start) :]

            before_el = _build_text_run_element(p, before_text, font_name, font_size)
            if before_el is not None:
                paragraph_el.append(before_el)

            if not badge_inserted:
                paragraph_el.append(
                    _build_badge_run_element(
                        p,
                        img_path,
                        badge_width_cm,
                        badge_tag=badge_tag,
                    )
                )
                badge_inserted = True

            after_el = _build_text_run_element(p, after_text, font_name, font_size)
            if after_el is not None:
                paragraph_el.append(after_el)

        text_pos = child_end

    if not badge_inserted:
        paragraph_el.append(
            _build_badge_run_element(
                p,
                img_path,
                badge_width_cm,
                badge_tag=badge_tag,
            )
        )


def replace_question_prefix_marker_in_paragraph(
    p,
    markers: dict[str, str],
    badge_width_cm: float,
    font_name: str,
    font_size: int,
    badge_tag: str = BADGE_TAG,
) -> bool:
    text = p.text or ""
    if not text:
        return False

    match = QUESTION_PREFIX_MARKER_RE.match(text)
    if not match:
        return False

    markers_expanded = _expand_markers(markers)
    marker_group_name = "marker_a" if match.group("marker_a") is not None else "marker_b"
    marker_raw = match.group(marker_group_name)
    marker_key = _resolve_marker_key(marker_raw, markers_expanded)
    if marker_key is None:
        return False

    img_path = resolve_path(markers_expanded[marker_key])
    if not img_path.exists():
        raise FileNotFoundError(f"Imagem não encontrada para marcador '{marker_key}': {img_path}")

    marker_start = match.start(marker_group_name)
    marker_end = match.end(marker_group_name)
    prefix = text[:marker_start]
    suffix = text[marker_end:]

    # `p.clear()` remove OMML/drawings; aqui substituímos só o marcador textual.
    if paragraph_has_office_math(p) or paragraph_has_drawing(p):
        _replace_prefix_marker_preserving_non_text(
            p,
            marker_start=marker_start,
            marker_end=marker_end,
            img_path=img_path,
            badge_width_cm=badge_width_cm,
            font_name=font_name,
            font_size=font_size,
            badge_tag=badge_tag,
        )
        return True

    p.clear()

    if prefix:
        run_prefix = p.add_run(prefix)
        apply_run_font(run_prefix, font_name, font_size)

    insert_badge_run(p, img_path, badge_width_cm, badge_tag=badge_tag)

    if suffix:
        run_suffix = p.add_run(suffix)
        apply_run_font(run_suffix, font_name, font_size)

    return True


def replace_markers_in_paragraph(
    p,
    markers: dict[str, str],
    badge_width_cm: float,
    font_name: str,
    font_size: int,
    badge_tag: str = BADGE_TAG,
) -> bool:
    text = p.text or ""
    if not text:
        return False

    # Expande marcadores para variantes NFC/NFD e prioriza match mais longo.
    markers_expanded = _expand_markers(markers)
    keys = sorted(markers_expanded.keys(), key=len, reverse=True)

    if not any(k in text for k in keys):
        return False

    original = text
    p.clear()

    i = 0
    changed = False

    while i < len(original):
        found = None
        for k in keys:
            if original.startswith(k, i):
                found = k
                break

        if found is not None:
            img_path = resolve_path(markers_expanded[found])
            if not img_path.exists():
                raise FileNotFoundError(
                    f"Imagem não encontrada para marcador '{found}': {img_path}"
                )

            insert_badge_run(p, img_path, badge_width_cm, badge_tag=badge_tag)
            changed = True
            i += len(found)
            continue

        run = p.add_run(original[i])
        apply_run_font(run, font_name, font_size)
        i += 1

    return changed
