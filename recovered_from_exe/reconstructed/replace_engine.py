from __future__ import annotations

from pathlib import Path
import hashlib
import os
import re
import subprocess
import sys
import unicodedata
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

GABARITO_RE = re.compile(r"^\s*GABARITO:\s*[A-E]\s*$", re.IGNORECASE)
QUESTION_DIFFICULTY_RE = re.compile(
    r"^\s*(?:(?P<num>\d+)\s*[\.\)\-:]?\s*)?(?:NIVEL\s*[:\-]?\s*)?\(?\s*(?P<level>FACIL|MEDIA|DIFICIL)\s*\)?\s*[:\-\.]?\s*$"
)
BADGE_TAG = "BADGE_REPLACE_DOCX"
SECTION_TAG = "SECTION_BANNER_REPLACE_DOCX"


def _runtime_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


ENGINE_LOG = _runtime_dir() / "engine_debug.log"


def _elog(msg: str) -> None:
    try:
        with open(ENGINE_LOG, "a", encoding="utf-8") as f:
            f.write(msg.rstrip() + "\n")
    except Exception:
        return


def _assets_dir() -> Path:
    # Portabilidade: mantém Windows e adiciona caminho padrão para macOS/Linux.
    if sys.platform == "darwin":
        base = Path.home() / "Library" / "Application Support"
    elif os.name == "nt":
        base = Path(
            os.environ.get("APPDATA")
            or (Path.home() / "AppData" / "Roaming")
        )
    else:
        base = Path(
            os.environ.get("XDG_DATA_HOME")
            or (Path.home() / ".local" / "share")
        )
    return base / "ReplaceDocx" / "assets"


def _normalize_area_slug(area: str) -> str:
    txt = unicodedata.normalize("NFD", (area or "").strip().lower())
    txt = "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")
    txt = re.sub(r"[^a-z0-9]+", "_", txt).strip("_")
    return txt or "geral"


def _normalize_text_key(text: str) -> str:
    txt = unicodedata.normalize("NFD", (text or "").strip().upper())
    txt = "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")
    txt = re.sub(r"\s+", " ", txt)
    return txt


def _default_markers_for_area(area: str) -> dict[str, str]:
    area_slug = _normalize_area_slug(area)
    base = f"areas/{area_slug}/capsulas"
    facil = f"{base}/facil.png"
    media = f"{base}/media.png"
    dificil = f"{base}/dificil.png"
    return {
        "(FÁCIL)": facil,
        "(MÉDIA)": media,
        "(DIFÍCIL)": dificil,
        "FÁCIL": facil,
        "MÉDIA": media,
        "DIFÍCIL": dificil,
        "FACIL": facil,
        "MEDIA": media,
        "DIFICIL": dificil,
    }


def _default_section_banners_for_area(area: str) -> dict[str, str]:
    area_slug = _normalize_area_slug(area)
    base = f"areas/{area_slug}/secoes"
    return {
        "EXERCÍCIOS DE SALA": f"{base}/exercicios_sala.png",
        "EXERCÍCIOS PROPOSTOS": f"{base}/exercicios_propostos.png",
        "SEÇÃO ENEM": f"{base}/secao_enem.png",
        "EXERCÍCIOS DE APROFUNDAMENTO": f"{base}/exercicios_aprofundamento.png",
        "EXERCÍCIOS REGIONAIS": f"{base}/exercicios_regionais.png",
        "EXERCÍCIO DISSERTATIVO": f"{base}/exercicios_dissertativos.png",
        "EXERCÍCIOS DISSERTATIVOS": f"{base}/exercicios_dissertativos.png",
    }


def _section_aliases_for_report() -> dict[str, list[str]]:
    return {
        "EXERCÍCIOS DE SALA": ["EXERCÍCIOS DE SALA"],
        "EXERCÍCIOS PROPOSTOS": ["EXERCÍCIOS PROPOSTOS"],
        "SEÇÃO ENEM": ["SEÇÃO ENEM"],
        "EXERCÍCIOS DE APROFUNDAMENTO": ["EXERCÍCIOS DE APROFUNDAMENTO"],
        "EXERCÍCIOS REGIONAIS": ["EXERCÍCIOS REGIONAIS"],
        "EXERCÍCIOS DISSERTATIVOS": ["EXERCÍCIOS DISSERTATIVOS", "EXERCÍCIO DISSERTATIVO"],
    }


def _match_section_for_report(norm_txt: str) -> str | None:
    aliases = _section_aliases_for_report()
    for canonical, raw_aliases in aliases.items():
        for alias in raw_aliases:
            norm_alias = _normalize_text_key(alias)
            if norm_txt == norm_alias:
                return canonical
            if norm_txt.startswith(norm_alias + ":"):
                return canonical
            if norm_txt.startswith(norm_alias + " -"):
                return canonical
            if norm_alias in norm_txt and len(norm_txt) <= len(norm_alias) + 20:
                return canonical
    return None


def _extract_difficulty(text: str) -> str | None:
    norm_txt = _normalize_text_key(text)
    match = QUESTION_DIFFICULTY_RE.match(norm_txt)
    if not match:
        return None
    raw = match.group("level")
    if raw == "FACIL":
        return "facil"
    if raw == "MEDIA":
        return "media"
    if raw == "DIFICIL":
        return "dificil"
    return None


def _extract_question_info(text: str) -> tuple[str | None, str] | None:
    norm_txt = _normalize_text_key(text)
    match = QUESTION_DIFFICULTY_RE.match(norm_txt)
    if not match:
        return None
    num = match.group("num")
    level = match.group("level")
    if level == "FACIL":
        return num, "facil"
    if level == "MEDIA":
        return num, "media"
    if level == "DIFICIL":
        return num, "dificil"
    return None


def _difficulty_label(diff_key: str) -> str:
    if diff_key == "facil":
        return "Fácil"
    if diff_key == "media":
        return "Média"
    if diff_key == "dificil":
        return "Difícil"
    return diff_key


def _set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def _set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def _style_cell_text(
    cell,
    *,
    align=WD_PARAGRAPH_ALIGNMENT.LEFT,
    bold: bool = False,
    color_hex: str | None = None,
    size_pt: int = 10,
) -> None:
    if not cell.paragraphs:
        return
    p = cell.paragraphs[0]
    p.alignment = align
    if not p.runs:
        p.add_run("")
    for run in p.runs:
        run.bold = bold
        run.font.size = Pt(size_pt)
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(color_hex)


def _collect_questions_by_section(doc: Document) -> list[dict]:
    sections: list[dict] = []
    current: dict | None = None
    seq = 0

    for p in list(doc.paragraphs):
        txt = (p.text or "").strip()
        if not txt:
            continue

        norm_txt = _normalize_text_key(txt)
        section = _match_section_for_report(norm_txt)
        if section:
            current = {
                "secao": section,
                "questoes": [],
            }
            sections.append(current)
            seq = 0
            continue

        if current is None:
            continue

        info = _extract_question_info(txt)
        if not info:
            continue

        num, diff = info
        seq += 1
        q_label = f"Questão {num}" if num else f"Questão {seq}"
        current["questoes"].append({"questao": q_label, "dificuldade": diff})

    return sections


def _insert_question_difficulty_table(
    doc: Document,
    section_item: dict,
    column_width_cm: float,
) -> int:
    questoes = section_item.get("questoes") or []
    if not questoes:
        return 0

    secao = section_item["secao"]
    title_p = doc.add_paragraph(f"Autoavaliação - {secao}")

    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if title_p.runs:
        title_p.runs[0].bold = True
        title_p.runs[0].font.size = Pt(11)
        title_p.runs[0].font.color.rgb = RGBColor(31, 41, 55)
    fmt = title_p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    table = doc.add_table(rows=2 + len(questoes), cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = False
    table_width_cm = max(5.0, float(column_width_cm))
    _set_table_width(table, table_width_cm)

    # Proporções: Questão | Dificuldade | Acertou | Errou | Revisar
    col_ratios = [0.30, 0.24, 0.15, 0.15, 0.16]
    col_widths = [table_width_cm * ratio for ratio in col_ratios]
    for row in table.rows:
        for idx, width_cm in enumerate(col_widths):
            _set_cell_width(row.cells[idx], width_cm)

    headers = ["Questão", "Dificuldade", "Acertou", "Errou", "Revisar"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        _set_cell_fill(cell, "1F2937")
        _style_cell_text(
            cell,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            color_hex="FFFFFF",
            size_pt=10,
        )

    counts = {"facil": 0, "media": 0, "dificil": 0}
    for i, q in enumerate(questoes, start=1):
        diff = q["dificuldade"]
        counts[diff] += 1
        row = table.rows[i]
        row.cells[0].text = q["questao"]
        row.cells[1].text = _difficulty_label(diff)
        row.cells[2].text = "☐"
        row.cells[3].text = "☐"
        row.cells[4].text = "☐"

        # Zebra rows para leitura.
        if i % 2 == 0:
            for c in row.cells:
                _set_cell_fill(c, "F8FAFC")

        diff_fill = {
            "facil": "DCFCE7",
            "media": "FEF3C7",
            "dificil": "FEE2E2",
        }.get(diff, "EEF2FF")
        _set_cell_fill(row.cells[1], diff_fill)

        _style_cell_text(row.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, size_pt=10)
        _style_cell_text(row.cells[1], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, size_pt=10)
        for col in (2, 3, 4):
            _style_cell_text(
                row.cells[col],
                align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                bold=True,
                color_hex="374151",
                size_pt=12,
            )

    # Linha de resumo visual no rodapé da tabela.
    summary_idx = len(questoes) + 1
    srow = table.rows[summary_idx]
    srow.cells[0].text = "Resumo"
    srow.cells[1].text = (
        f"Fácil {counts['facil']} | Média {counts['media']} | Difícil {counts['dificil']}"
    )
    srow.cells[2].text = "Acertos: ____"
    srow.cells[3].text = "Erros: ____"
    srow.cells[4].text = "Revisar: ____"
    for c in srow.cells:
        _set_cell_fill(c, "E5E7EB")
    _style_cell_text(srow.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
    _style_cell_text(srow.cells[1], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
    _style_cell_text(srow.cells[2], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
    _style_cell_text(srow.cells[3], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
    _style_cell_text(srow.cells[4], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

    # Espaço entre tabelas no bloco final de autoavaliação.
    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)

    return 1


def _insert_question_difficulty_tables(
    doc: Document,
    sections_data: list[dict],
    column_width_cm: float,
) -> int:
    inserted = 0
    if not sections_data:
        return inserted

    # Todas as autoavaliações ficam no final do documento, na ordem das seções.
    block_title = doc.add_paragraph("Quadro de Autoavaliação por Seção")
    block_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if block_title.runs:
        block_title.runs[0].bold = True
        block_title.runs[0].font.size = Pt(12)
        block_title.runs[0].font.color.rgb = RGBColor(17, 24, 39)
    bfmt = block_title.paragraph_format
    bfmt.space_before = Pt(10)
    bfmt.space_after = Pt(8)
    bfmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for item in sections_data:
        inserted += _insert_question_difficulty_table(
            doc,
            item,
            column_width_cm=column_width_cm,
        )
    return inserted


def _append_difficulty_report_appendix(doc: Document, report: dict) -> bool:
    sections = report.get("secoes") if isinstance(report, dict) else None
    if not sections:
        return False

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    try:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
    except Exception:
        pass

    try:
        text_width_cm = (
            float(section.page_width.cm)
            - float(section.left_margin.cm)
            - float(section.right_margin.cm)
        )
    except Exception:
        text_width_cm = 17.4
    text_width_cm = max(10.0, text_width_cm)

    title = doc.add_paragraph("Relatório de Dificuldade por Seção")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if title.runs:
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(13)
        title.runs[0].font.color.rgb = RGBColor(17, 24, 39)
    tfmt = title.paragraph_format
    tfmt.space_before = Pt(0)
    tfmt.space_after = Pt(6)

    meta = [
        f"Conteúdo: {report.get('conteudo', '-')}",
        f"Área: {report.get('area_conhecimento', '-')}",
        f"Gerado em: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    for line in meta:
        p = doc.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(75, 85, 99)

    totals = report.get("totais") or {"facil": 0, "media": 0, "dificil": 0, "total": 0}
    cards = doc.add_table(rows=2, cols=4)
    cards.autofit = False
    _set_table_width(cards, text_width_cm)
    card_width_cm = text_width_cm / 4
    for row in cards.rows:
        for c in row.cells:
            _set_cell_width(c, card_width_cm)
    card_headers = ["Fácil", "Média", "Difícil", "Total"]
    card_values = [totals["facil"], totals["media"], totals["dificil"], totals["total"]]
    header_fill = ["DCFCE7", "FEF3C7", "FEE2E2", "E5E7EB"]
    value_color = ["166534", "92400E", "991B1B", "111827"]
    for idx in range(4):
        cards.cell(0, idx).text = card_headers[idx]
        _set_cell_fill(cards.cell(0, idx), header_fill[idx])
        _style_cell_text(
            cards.cell(0, idx),
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            size_pt=9,
        )
        cards.cell(1, idx).text = str(card_values[idx])
        _style_cell_text(
            cards.cell(1, idx),
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            color_hex=value_color[idx],
            size_pt=12,
        )

    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=2 + len(sections), cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = False
    _set_table_width(table, text_width_cm)

    col_ratios = [0.52, 0.12, 0.12, 0.12, 0.12]
    col_widths = [text_width_cm * ratio for ratio in col_ratios]
    for row in table.rows:
        for idx, width_cm in enumerate(col_widths):
            _set_cell_width(row.cells[idx], width_cm)

    headers = ["Seção", "Fácil", "Média", "Difícil", "Total"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        _set_cell_fill(cell, "1F2937")
        _style_cell_text(
            cell,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            color_hex="FFFFFF",
            size_pt=9,
        )

    for idx, row_data in enumerate(sections, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(row_data["secao"])
        row.cells[1].text = str(row_data["facil"])
        row.cells[2].text = str(row_data["media"])
        row.cells[3].text = str(row_data["dificil"])
        row.cells[4].text = str(row_data["total"])

        if idx % 2 == 0:
            for c in row.cells:
                _set_cell_fill(c, "F8FAFC")

        _style_cell_text(row.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True, size_pt=9)
        _style_cell_text(row.cells[1], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="166534", size_pt=9)
        _style_cell_text(row.cells[2], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="92400E", size_pt=9)
        _style_cell_text(row.cells[3], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="991B1B", size_pt=9)
        _style_cell_text(row.cells[4], align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, color_hex="111827", size_pt=9)

    total_row_idx = len(sections) + 1
    total_row = table.rows[total_row_idx]
    total_row.cells[0].text = "TOTAL"
    total_row.cells[1].text = str(totals["facil"])
    total_row.cells[2].text = str(totals["media"])
    total_row.cells[3].text = str(totals["dificil"])
    total_row.cells[4].text = str(totals["total"])
    for c in total_row.cells:
        _set_cell_fill(c, "E5E7EB")
        _style_cell_text(
            c,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            color_hex="111827",
            size_pt=9,
        )
    _style_cell_text(total_row.cells[0], align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True, color_hex="111827", size_pt=9)

    return True


def gerar_relatorio_dificuldade_por_secao(
    input_path: str | Path,
    area_conhecimento: str = "biologia",
) -> dict:
    input_path = Path(input_path)
    doc = Document(str(input_path))

    section_order = list(_section_aliases_for_report().keys())
    section_counts: dict[str, dict[str, int]] = {
        section: {"facil": 0, "media": 0, "dificil": 0}
        for section in section_order
    }
    section_counts["SEM SEÇÃO"] = {"facil": 0, "media": 0, "dificil": 0}

    current_section = "SEM SEÇÃO"
    for p in _iter_paragraphs(doc):
        txt = (p.text or "").strip()
        if not txt:
            continue

        norm_txt = _normalize_text_key(txt)
        section = _match_section_for_report(norm_txt)
        if section:
            current_section = section
            continue

        difficulty = _extract_difficulty(txt)
        if not difficulty:
            continue

        section_counts[current_section][difficulty] += 1

    sections = []
    for section in section_order + ["SEM SEÇÃO"]:
        counts = section_counts[section]
        total = counts["facil"] + counts["media"] + counts["dificil"]
        if total == 0 and section == "SEM SEÇÃO":
            continue
        sections.append(
            {
                "secao": section,
                "facil": counts["facil"],
                "media": counts["media"],
                "dificil": counts["dificil"],
                "total": total,
            }
        )

    total_facil = sum(item["facil"] for item in sections)
    total_media = sum(item["media"] for item in sections)
    total_dificil = sum(item["dificil"] for item in sections)

    return {
        "conteudo": input_path.stem,
        "arquivo_origem": str(input_path),
        "area_conhecimento": area_conhecimento,
        "secoes": sections,
        "totais": {
            "facil": total_facil,
            "media": total_media,
            "dificil": total_dificil,
            "total": total_facil + total_media + total_dificil,
        },
    }


def _paths_with_image_extension_fallback(p: Path) -> list[Path]:
    # Prioridade: JPG/JPEG primeiro para refletir artes novas em JPG.
    preferred_exts = [".jpg", ".jpeg", ".png"]
    image_exts = {".jpg", ".jpeg", ".png"}

    if p.suffix.lower() in image_exts:
        stem = p.with_suffix("")
        return [stem.with_suffix(ext) for ext in preferred_exts]

    if p.suffix:
        return [p]

    return [p.with_suffix(ext) for ext in preferred_exts]


def _convert_image_for_docx(img_path: Path) -> Path:
    if img_path.suffix.lower() not in {".jpg", ".jpeg"}:
        return img_path

    cache_dir = _runtime_dir() / ".image_cache_docx"
    cache_dir.mkdir(parents=True, exist_ok=True)
    key = f"{img_path.resolve()}::{img_path.stat().st_mtime_ns}"
    out_name = hashlib.sha1(key.encode("utf-8")).hexdigest() + ".png"
    out_path = cache_dir / out_name
    if out_path.exists():
        return out_path

    # Tenta Pillow primeiro (se existir no ambiente).
    try:
        from PIL import Image  # type: ignore

        with Image.open(str(img_path)) as im:
            if im.mode not in {"RGB", "RGBA"}:
                im = im.convert("RGB")
            im.save(str(out_path), format="PNG")
        if out_path.exists():
            return out_path
    except Exception:
        pass

    # Fallback macOS via sips.
    if sys.platform == "darwin":
        try:
            subprocess.run(
                ["sips", "-s", "format", "png", str(img_path), "--out", str(out_path)],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            if out_path.exists():
                return out_path
        except Exception:
            pass

    return img_path


def _add_picture_resilient(run, img_path: Path, width_cm: float):
    try:
        return run.add_picture(str(img_path), width=Cm(width_cm))
    except UnrecognizedImageError:
        converted = _convert_image_for_docx(img_path)
        if converted != img_path and converted.exists():
            return run.add_picture(str(converted), width=Cm(width_cm))
        raise


def _resolve_path(p: str | Path) -> Path:
    p = Path(p)

    if p.is_absolute() and p.exists():
        return p

    assets = _assets_dir()
    base = _runtime_dir()
    project_assets = base.parent.parent / "Assets"
    candidates = [
        (assets / p),
        (assets / p.name),
        (base / p),
        (base / "assets" / p),
        (base.parent / "Assets" / p),
        (base.parent / "Assets" / p.name),
        (project_assets / p),
        (project_assets / p.name),
    ]
    expanded: list[Path] = []
    for cand in candidates:
        expanded.extend(_paths_with_image_extension_fallback(cand))

    seen = set()
    for cand in expanded:
        cand_resolved = cand.resolve()
        key = str(cand_resolved)
        if key in seen:
            continue
        seen.add(key)
        if cand_resolved.exists():
            return cand_resolved

    return (base / p).resolve()


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _remove_paragraph(p) -> None:
    p._element.getparent().remove(p._element)


def _apply_paragraph_layout(p, justify: bool) -> None:
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if justify
        else WD_PARAGRAPH_ALIGNMENT.LEFT
    )


def _apply_run_font(run, font_name: str, font_size: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)


def _insert_badge_run(p, img_path: Path, badge_width_cm: float) -> None:
    run = p.add_run()
    pic = _add_picture_resilient(run, img_path, badge_width_cm)

    # Marca a imagem para o backend do Word ignorar badges na etapa final.
    try:
        pic._inline.docPr.set("descr", BADGE_TAG)
        pic._inline.docPr.set("title", BADGE_TAG)
    except Exception:
        return


def _replace_markers_in_paragraph(
    p,
    markers: dict[str, str],
    badge_width_cm: float,
    font_name: str,
    font_size: int,
) -> bool:
    text = p.text or ""
    if not text:
        return False

    # Expande marcadores para variantes NFC/NFD e prioriza match mais longo.
    markers_expanded: dict[str, str] = {}
    for raw_key, img in markers.items():
        for variant in (
            raw_key,
            unicodedata.normalize("NFC", raw_key),
            unicodedata.normalize("NFD", raw_key),
        ):
            if variant and variant not in markers_expanded:
                markers_expanded[variant] = img

    keys = sorted(markers_expanded.keys(), key=len, reverse=True)

    if not any(k in text for k in keys):
        return False

    original = text
    p.clear()

    i = 0
    changed = False

    while i < len(original):
        found = None
        for k in keys:
            if original.startswith(k, i):
                found = k
                break

        if found is not None:
            img_path = _resolve_path(markers_expanded[found])
            if not img_path.exists():
                raise FileNotFoundError(
                    f"Imagem não encontrada para marcador '{found}': {img_path}"
                )

            _insert_badge_run(p, img_path, badge_width_cm)
            changed = True
            i += len(found)
            continue

        run = p.add_run(original[i])
        _apply_run_font(run, font_name, font_size)
        i += 1

    return changed


def _safe_tag_suffix(text: str) -> str:
    key = _normalize_text_key(text)
    return re.sub(r"[^A-Z0-9]+", "_", key).strip("_") or "GENERIC"


def _replace_paragraph_with_section_banner(
    p,
    img_path: Path,
    section_banner_width_cm: float,
    section_tag: str,
) -> None:
    # Substitui totalmente o título da seção pelo banner.
    p.clear()
    banner_run = p.add_run()
    pic = _add_picture_resilient(banner_run, img_path, section_banner_width_cm)

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    try:
        pic._inline.docPr.set("descr", section_tag)
        pic._inline.docPr.set("title", section_tag)
    except Exception:
        pass


def _apply_section_banners(
    doc: Document,
    section_banners: dict[str, str],
    section_banner_width_cm: float,
) -> int:
    normalized_map: dict[str, str] = {}
    for title, img in section_banners.items():
        norm_title = _normalize_text_key(title)
        if norm_title:
            normalized_map[norm_title] = img

    inserted = 0
    for p in list(_iter_paragraphs(doc)):
        txt = (p.text or "").strip()
        if not txt:
            continue

        norm_txt = _normalize_text_key(txt)
        if not norm_txt or len(norm_txt) > 120:
            continue

        matched_title = None
        matched_img = None
        for norm_title, img in normalized_map.items():
            if (
                norm_txt == norm_title
                or norm_txt.startswith(norm_title + " ")
                or norm_title in norm_txt
            ):
                matched_title = norm_title
                matched_img = img
                break

        if matched_img is None:
            continue

        section_tag = f"{SECTION_TAG}_{_safe_tag_suffix(matched_title)}"

        img_path = _resolve_path(matched_img)
        if not img_path.exists():
            _elog(f"Section banner missing for '{txt}': {img_path}")
            continue

        try:
            _replace_paragraph_with_section_banner(
                p,
                img_path=img_path,
                section_banner_width_cm=section_banner_width_cm,
                section_tag=section_tag,
            )
            inserted += 1
        except Exception as exc:
            _elog(f"Section banner failed for '{txt}': {img_path} ({exc})")
            continue

    return inserted


def processar_docx(input_path: str | Path, output_path: str | Path, config: dict):
    _elog("\n============================")
    _elog("START " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {input_path}")

    font_name = (config.get("font_name") or "Arial").strip()
    font_size = int(config.get("font_size", 11))
    badge_width_cm = float(config.get("badge_width_cm", 1.3))
    column_width_cm = float(config.get("column_width_cm", 7.5))
    remove_gabarito = bool(config.get("remove_gabarito", True))
    justify = bool(config.get("justify", True))
    area_conhecimento = (config.get("area_conhecimento") or "biologia").strip()
    insert_section_banners = bool(config.get("insert_section_banners", True))
    insert_question_tables = bool(config.get("insert_question_tables", True))
    append_difficulty_report = bool(config.get("append_difficulty_report", True))
    difficulty_report_data = config.get("difficulty_report_data")
    if difficulty_report_data is not None and not isinstance(difficulty_report_data, dict):
        raise ValueError("Config inválida: 'difficulty_report_data' deve ser um dict.")
    section_banner_width_raw = config.get("section_banner_width_cm")
    section_banner_width_cm = float(
        section_banner_width_raw
        if section_banner_width_raw is not None
        else column_width_cm
    )

    base_markers = _default_markers_for_area(area_conhecimento)
    user_markers = config.get("markers") or {}
    if user_markers and not isinstance(user_markers, dict):
        raise ValueError("Config inválida: 'markers' deve ser um dict.")
    markers = dict(base_markers)
    markers.update(user_markers)
    if not markers:
        raise ValueError("Config inválida: 'markers' deve ser um dict com pelo menos 1 item.")

    base_sections = _default_section_banners_for_area(area_conhecimento)
    user_sections = config.get("section_banners") or {}
    if user_sections and not isinstance(user_sections, dict):
        raise ValueError("Config inválida: 'section_banners' deve ser um dict.")
    section_banners = dict(base_sections)
    section_banners.update(user_sections)

    finalize_word = bool(config.get("finalize_word", True))
    force_inline_wrap = bool(config.get("force_inline_wrap", True))

    doc = Document(str(input_path))

    for p in _iter_paragraphs(doc):
        _apply_paragraph_layout(p, justify)
        for r in p.runs:
            _apply_run_font(r, font_name, font_size)

    if remove_gabarito:
        to_delete = []
        for p in _iter_paragraphs(doc):
            txt = (p.text or "").strip()
            if GABARITO_RE.match(txt):
                to_delete.append(p)

        for p in to_delete:
            _remove_paragraph(p)

        _elog("Removed gabarito lines: " + str(len(to_delete)))

    sections_data = _collect_questions_by_section(doc) if insert_question_tables else []

    if insert_section_banners and section_banners:
        inserted = _apply_section_banners(
            doc,
            section_banners=section_banners,
            section_banner_width_cm=section_banner_width_cm,
        )
        _elog("Inserted section banners: " + str(inserted))

    for p in _iter_paragraphs(doc):
        _replace_markers_in_paragraph(
            p, markers, badge_width_cm, font_name, font_size
        )

    if insert_question_tables and sections_data:
        inserted_tables = _insert_question_difficulty_tables(
            doc,
            sections_data,
            column_width_cm=column_width_cm,
        )
        _elog("Inserted question difficulty tables: " + str(inserted_tables))

    if append_difficulty_report:
        report_data = difficulty_report_data or gerar_relatorio_dificuldade_por_secao(
            input_path,
            area_conhecimento=area_conhecimento,
        )
        if _append_difficulty_report_appendix(doc, report_data):
            _elog("Inserted A4 difficulty report appendix at document end.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    _elog("Saved via python-docx OK")

    if finalize_word:
        from word_finalize import finalize_with_word

        finalize_with_word(
            docx_in=output_path,
            docx_out=output_path,
            column_width_cm=column_width_cm,
            font_name=font_name,
            font_size=font_size,
            justify=justify,
            force_inline_wrap=force_inline_wrap,
            badge_tag=BADGE_TAG,
        )
        _elog("Finalized with Word OK")

    return {"arquivo_saida": str(output_path)}
