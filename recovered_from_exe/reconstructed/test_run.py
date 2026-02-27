from __future__ import annotations

from pathlib import Path

from docx import Document

from replace_engine import processar_docx


def build_input(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Questao 1 (FÁCIL)")
    doc.add_paragraph("Texto da questao em duas linhas para validar formatacao.")
    doc.add_paragraph("GABARITO: A")
    doc.add_paragraph("Questao 2 (MÉDIA)")
    doc.add_paragraph("Mais texto aqui.")
    doc.add_paragraph("Questao 3 (DIFÍCIL)")
    doc.save(str(path))


def main() -> None:
    base = Path(__file__).resolve().parent
    entrada = base / "entrada_teste.docx"
    saida = base / "saida_teste.docx"

    build_input(entrada)

    cfg = {
        "area_conhecimento": "biologia",
        "font_name": "Arial",
        "font_size": 11,
        "badge_width_cm": 1.3,
        "column_width_cm": 7.5,
        "remove_gabarito": True,
        "justify": True,
        "finalize_word": False,
        "force_inline_wrap": True,
    }

    result = processar_docx(entrada, saida, cfg)
    print("OK", result)
    print("entrada:", entrada)
    print("saida:", saida)
    print("log:", base / "engine_debug.log")


if __name__ == "__main__":
    main()
