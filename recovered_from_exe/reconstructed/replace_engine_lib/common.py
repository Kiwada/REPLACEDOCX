from __future__ import annotations

import hashlib
import os
import re
import subprocess
import sys
import unicodedata
from pathlib import Path

from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Cm

GABARITO_RE = re.compile(r"^\s*GABARITO:\s*[A-E]\s*$", re.IGNORECASE)
QUESTION_DIFFICULTY_RE = re.compile(
    r"^\s*(?:(?P<num>\d+)\s*[\.\)\-:]?\s*)?(?:NIVEL\s*[:\-]?\s*)?\(?\s*(?P<level>FACIL|MEDIA|DIFICIL)\s*\)?\s*[:\-\.]?\s*$"
)
BADGE_TAG = "BADGE_REPLACE_DOCX"
SECTION_TAG = "SECTION_BANNER_REPLACE_DOCX"


def runtime_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    # common.py está em .../reconstructed/replace_engine_lib/common.py
    # runtime da engine deve seguir em .../reconstructed
    return Path(__file__).resolve().parents[1]


ENGINE_LOG = runtime_dir() / "engine_debug.log"


def elog(msg: str) -> None:
    try:
        with open(ENGINE_LOG, "a", encoding="utf-8") as f:
            f.write(msg.rstrip() + "\n")
    except Exception:
        return


def assets_dir() -> Path:
    # Portabilidade: mantém Windows e adiciona caminho padrão para macOS/Linux.
    if sys.platform == "darwin":
        base = Path.home() / "Library" / "Application Support"
    elif os.name == "nt":
        base = Path(os.environ.get("APPDATA") or (Path.home() / "AppData" / "Roaming"))
    else:
        base = Path(os.environ.get("XDG_DATA_HOME") or (Path.home() / ".local" / "share"))
    return base / "ReplaceDocx" / "assets"


def normalize_area_slug(area: str) -> str:
    txt = unicodedata.normalize("NFD", (area or "").strip().lower())
    txt = "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")
    txt = re.sub(r"[^a-z0-9]+", "_", txt).strip("_")
    return txt or "geral"


def normalize_text_key(text: str) -> str:
    txt = unicodedata.normalize("NFD", (text or "").strip().upper())
    txt = "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")
    txt = re.sub(r"\s+", " ", txt)
    return txt


def safe_tag_suffix(text: str) -> str:
    key = normalize_text_key(text)
    return re.sub(r"[^A-Z0-9]+", "_", key).strip("_") or "GENERIC"


def default_markers_for_area(area: str) -> dict[str, str]:
    area_slug = normalize_area_slug(area)
    base = f"areas/{area_slug}/capsulas"
    facil = f"{base}/facil.png"
    media = f"{base}/media.png"
    dificil = f"{base}/dificil.png"
    return {
        "(FÁCIL)": facil,
        "(MÉDIA)": media,
        "(DIFÍCIL)": dificil,
        "FÁCIL": facil,
        "MÉDIA": media,
        "DIFÍCIL": dificil,
        "FACIL": facil,
        "MEDIA": media,
        "DIFICIL": dificil,
    }


def default_section_banners_for_area(area: str) -> dict[str, str]:
    area_slug = normalize_area_slug(area)
    base = f"areas/{area_slug}/secoes"
    return {
        "EXERCÍCIOS DE SALA": f"{base}/exercicios_sala.png",
        "EXERCÍCIOS PROPOSTOS": f"{base}/exercicios_propostos.png",
        "SEÇÃO ENEM": f"{base}/secao_enem.png",
        "EXERCÍCIOS DE APROFUNDAMENTO": f"{base}/exercicios_aprofundamento.png",
        "EXERCÍCIOS REGIONAIS": f"{base}/exercicios_regionais.png",
        "EXERCÍCIO DISSERTATIVO": f"{base}/exercicios_dissertativos.png",
        "EXERCÍCIOS DISSERTATIVOS": f"{base}/exercicios_dissertativos.png",
    }


def paths_with_image_extension_fallback(p: Path) -> list[Path]:
    # Prioridade: JPG/JPEG primeiro para refletir artes novas em JPG.
    preferred_exts = [".jpg", ".jpeg", ".png"]
    image_exts = {".jpg", ".jpeg", ".png"}

    if p.suffix.lower() in image_exts:
        stem = p.with_suffix("")
        return [stem.with_suffix(ext) for ext in preferred_exts]

    if p.suffix:
        return [p]

    return [p.with_suffix(ext) for ext in preferred_exts]


def convert_image_for_docx(img_path: Path) -> Path:
    if img_path.suffix.lower() not in {".jpg", ".jpeg"}:
        return img_path

    cache_dir = runtime_dir() / ".image_cache_docx"
    cache_dir.mkdir(parents=True, exist_ok=True)
    key = f"{img_path.resolve()}::{img_path.stat().st_mtime_ns}"
    out_name = hashlib.sha1(key.encode("utf-8")).hexdigest() + ".png"
    out_path = cache_dir / out_name
    if out_path.exists():
        return out_path

    # Tenta Pillow primeiro (se existir no ambiente).
    try:
        from PIL import Image  # type: ignore

        with Image.open(str(img_path)) as im:
            if im.mode not in {"RGB", "RGBA"}:
                im = im.convert("RGB")
            im.save(str(out_path), format="PNG")
        if out_path.exists():
            return out_path
    except Exception:
        pass

    # Fallback macOS via sips.
    if sys.platform == "darwin":
        try:
            subprocess.run(
                ["sips", "-s", "format", "png", str(img_path), "--out", str(out_path)],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            if out_path.exists():
                return out_path
        except Exception:
            pass

    return img_path


def add_picture_resilient(run, img_path: Path, width_cm: float):
    try:
        return run.add_picture(str(img_path), width=Cm(width_cm))
    except UnrecognizedImageError:
        converted = convert_image_for_docx(img_path)
        if converted != img_path and converted.exists():
            return run.add_picture(str(converted), width=Cm(width_cm))
        raise


def resolve_path(p: str | Path) -> Path:
    p = Path(p)

    if p.is_absolute() and p.exists():
        return p

    assets = assets_dir()
    base = runtime_dir()
    project_assets = base.parent.parent / "Assets"
    candidates = [
        (assets / p),
        (assets / p.name),
        (base / p),
        (base / "assets" / p),
        (base.parent / "Assets" / p),
        (base.parent / "Assets" / p.name),
        (project_assets / p),
        (project_assets / p.name),
    ]
    expanded: list[Path] = []
    for cand in candidates:
        expanded.extend(paths_with_image_extension_fallback(cand))

    seen = set()
    for cand in expanded:
        cand_resolved = cand.resolve()
        key = str(cand_resolved)
        if key in seen:
            continue
        seen.add(key)
        if cand_resolved.exists():
            return cand_resolved

    return (base / p).resolve()
