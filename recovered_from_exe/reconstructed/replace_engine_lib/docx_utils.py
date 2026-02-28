from __future__ import annotations

import re
import unicodedata

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .common import BADGE_TAG, add_picture_resilient, resolve_path


def set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def style_cell_text(
    cell,
    *,
    align=WD_PARAGRAPH_ALIGNMENT.LEFT,
    bold: bool = False,
    color_hex: str | None = None,
    size_pt: int = 10,
) -> None:
    if not cell.paragraphs:
        return
    p = cell.paragraphs[0]
    p.alignment = align
    if not p.runs:
        p.add_run("")
    for run in p.runs:
        run.bold = bold
        run.font.size = Pt(size_pt)
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(color_hex)


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def remove_paragraph(p) -> None:
    p._element.getparent().remove(p._element)


def apply_paragraph_layout(p, justify: bool) -> None:
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY if justify else WD_PARAGRAPH_ALIGNMENT.LEFT


def apply_run_font(run, font_name: str, font_size: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)


def is_reference_caption_text(text: str) -> bool:
    txt = (text or "").strip()
    if not txt:
        return False
    norm = unicodedata.normalize("NFD", txt.upper())
    norm = "".join(ch for ch in norm if unicodedata.category(ch) != "Mn")
    norm = re.sub(r"\s+", " ", norm)
    return bool(re.match(r"^(FONTE|REFERENCIA|REFERENCIAS|DISPONIVEL\s+EM)\b", norm))


def is_figure_caption_text(text: str) -> bool:
    txt = (text or "").strip()
    if not txt:
        return False
    norm = unicodedata.normalize("NFD", txt.upper())
    norm = "".join(ch for ch in norm if unicodedata.category(ch) != "Mn")
    norm = re.sub(r"\s+", " ", norm)
    return bool(re.match(r"^(FIGURA|IMAGEM|ILUSTRACAO|FOTO|GRAFICO|QUADRO|TABELA)\b", norm))


def paragraph_has_drawing(paragraph) -> bool:
    try:
        return bool(paragraph._element.xpath(".//w:drawing"))
    except Exception:
        return False


def insert_badge_run(p, img_path, badge_width_cm: float, badge_tag: str = BADGE_TAG) -> None:
    run = p.add_run()
    pic = add_picture_resilient(run, img_path, badge_width_cm)

    # Marca a imagem para o backend do Word ignorar badges na etapa final.
    try:
        pic._inline.docPr.set("descr", badge_tag)
        pic._inline.docPr.set("title", badge_tag)
    except Exception:
        return


def replace_markers_in_paragraph(
    p,
    markers: dict[str, str],
    badge_width_cm: float,
    font_name: str,
    font_size: int,
    badge_tag: str = BADGE_TAG,
) -> bool:
    text = p.text or ""
    if not text:
        return False

    # Expande marcadores para variantes NFC/NFD e prioriza match mais longo.
    markers_expanded: dict[str, str] = {}
    for raw_key, img in markers.items():
        for variant in (
            raw_key,
            unicodedata.normalize("NFC", raw_key),
            unicodedata.normalize("NFD", raw_key),
        ):
            if variant and variant not in markers_expanded:
                markers_expanded[variant] = img

    keys = sorted(markers_expanded.keys(), key=len, reverse=True)

    if not any(k in text for k in keys):
        return False

    original = text
    p.clear()

    i = 0
    changed = False

    while i < len(original):
        found = None
        for k in keys:
            if original.startswith(k, i):
                found = k
                break

        if found is not None:
            img_path = resolve_path(markers_expanded[found])
            if not img_path.exists():
                raise FileNotFoundError(
                    f"Imagem não encontrada para marcador '{found}': {img_path}"
                )

            insert_badge_run(p, img_path, badge_width_cm, badge_tag=badge_tag)
            changed = True
            i += len(found)
            continue

        run = p.add_run(original[i])
        apply_run_font(run, font_name, font_size)
        i += 1

    return changed
