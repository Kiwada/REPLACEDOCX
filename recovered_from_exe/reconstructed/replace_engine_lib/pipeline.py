from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm

from .common import (
    ALTERNATIVE_ONLY_RE,
    BADGE_TAG,
    GABARITO_RE,
    default_markers_for_area,
    default_section_banners_for_area,
    elog,
)
from .docx_utils import (
    apply_paragraph_layout,
    apply_run_font,
    is_figure_caption_text,
    is_reference_caption_text,
    iter_paragraphs,
    paragraph_has_drawing,
    remove_paragraph,
    replace_markers_in_paragraph,
)
from .renderers import (
    append_difficulty_report_appendix,
    apply_section_banners,
    insert_question_difficulty_tables,
)
from .report import collect_questions_by_section, gerar_relatorio_dificuldade_por_secao


def _is_reference_below_image(paragraphs: list, idx: int, lookback: int = 3) -> bool:
    if idx < 0 or idx >= len(paragraphs):
        return False
    current = paragraphs[idx]
    if paragraph_has_drawing(current):
        return True

    lower = max(-1, idx - lookback - 1)
    for j in range(idx - 1, lower, -1):
        prev = paragraphs[j]
        if paragraph_has_drawing(prev):
            return True
        prev_text = (prev.text or "").strip()
        if prev_text and not (
            is_figure_caption_text(prev_text) or is_reference_caption_text(prev_text)
        ):
            return False
    return False


def _apply_reference_caption_font(doc: Document, font_name: str, reference_font_size: int = 8) -> int:
    updated = 0
    paragraph_groups = [list(doc.paragraphs)]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraph_groups.append(list(cell.paragraphs))

    for paragraphs in paragraph_groups:
        for i, p in enumerate(paragraphs):
            if not is_reference_caption_text(p.text or ""):
                continue
            if not _is_reference_below_image(paragraphs, i):
                continue
            for run in p.runs:
                apply_run_font(run, font_name, reference_font_size)
            updated += 1

    return updated


def processar_docx(input_path: str | Path, output_path: str | Path, config: dict):
    elog("\n============================")
    elog("START " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {input_path}")

    font_name = (config.get("font_name") or "Arial").strip()
    font_size = int(config.get("font_size", 11))
    badge_width_cm = float(config.get("badge_width_cm", 1.3))
    column_width_cm = float(config.get("column_width_cm", 7.5))
    remove_gabarito = bool(config.get("remove_gabarito", True))
    justify = bool(config.get("justify", True))
    format_text = bool(config.get("format_text", True))
    area_conhecimento = (config.get("area_conhecimento") or "biologia").strip()
    insert_section_banners = bool(config.get("insert_section_banners", True))
    insert_question_tables = bool(config.get("insert_question_tables", True))
    append_difficulty_report = bool(config.get("append_difficulty_report", True))
    difficulty_report_data = config.get("difficulty_report_data")
    if difficulty_report_data is not None and not isinstance(difficulty_report_data, dict):
        raise ValueError("Config inválida: 'difficulty_report_data' deve ser um dict.")
    margin_top_cm = float(config.get("margin_top_cm", 2.0))
    margin_bottom_cm = float(config.get("margin_bottom_cm", 2.0))
    margin_left_cm = float(config.get("margin_left_cm", 2.5))
    margin_right_cm = float(config.get("margin_right_cm", 2.5))
    section_banner_width_raw = config.get("section_banner_width_cm")
    section_banner_width_cm = float(
        section_banner_width_raw if section_banner_width_raw is not None else column_width_cm
    )

    base_markers = default_markers_for_area(area_conhecimento)
    user_markers = config.get("markers") or {}
    if user_markers and not isinstance(user_markers, dict):
        raise ValueError("Config inválida: 'markers' deve ser um dict.")
    markers = dict(base_markers)
    markers.update(user_markers)
    if not markers:
        raise ValueError("Config inválida: 'markers' deve ser um dict com pelo menos 1 item.")

    base_sections = default_section_banners_for_area(area_conhecimento)
    user_sections = config.get("section_banners") or {}
    if user_sections and not isinstance(user_sections, dict):
        raise ValueError("Config inválida: 'section_banners' deve ser um dict.")
    section_banners = dict(base_sections)
    section_banners.update(user_sections)

    finalize_word = bool(config.get("finalize_word", True))
    force_inline_wrap = bool(config.get("force_inline_wrap", True))

    doc = Document(str(input_path))
    for section in doc.sections:
        section.top_margin = Cm(margin_top_cm)
        section.bottom_margin = Cm(margin_bottom_cm)
        section.left_margin = Cm(margin_left_cm)
        section.right_margin = Cm(margin_right_cm)

    if format_text:
        for p in iter_paragraphs(doc):
            apply_paragraph_layout(p, justify)
            for r in p.runs:
                apply_run_font(r, font_name, font_size)
    else:
        elog("Skipped global text formatting (format_text=False).")

    if remove_gabarito:
        paragraphs = list(iter_paragraphs(doc))
        to_delete_idx: set[int] = set()

        for i, p in enumerate(paragraphs):
            txt = (p.text or "").strip()
            m = GABARITO_RE.match(txt)
            if not m:
                continue
            to_delete_idx.add(i)

            # Caso comum:
            # "Resposta:" em uma linha e alternativa ("A", "B", ...) na próxima.
            if m.groupdict().get("alt"):
                continue

            j = i + 1
            while j < len(paragraphs):
                next_txt = (paragraphs[j].text or "").strip()
                if not next_txt:
                    j += 1
                    continue
                if ALTERNATIVE_ONLY_RE.match(next_txt):
                    to_delete_idx.add(j)
                break

        for idx in sorted(to_delete_idx, reverse=True):
            remove_paragraph(paragraphs[idx])

        elog("Removed answer lines (gabarito/resposta): " + str(len(to_delete_idx)))

    sections_data = collect_questions_by_section(doc) if insert_question_tables else []

    if insert_section_banners and section_banners:
        inserted = apply_section_banners(
            doc,
            section_banners=section_banners,
            section_banner_width_cm=section_banner_width_cm,
        )
        elog("Inserted section banners: " + str(inserted))

    for p in iter_paragraphs(doc):
        replace_markers_in_paragraph(
            p,
            markers,
            badge_width_cm,
            font_name,
            font_size,
            badge_tag=BADGE_TAG,
        )

    ref_count = _apply_reference_caption_font(doc, font_name, reference_font_size=8)
    if ref_count:
        elog("Applied font size 8 to image references: " + str(ref_count))

    if insert_question_tables and sections_data:
        inserted_tables = insert_question_difficulty_tables(
            doc,
            sections_data,
            column_width_cm=column_width_cm,
            section_banners=section_banners,
            section_banner_width_cm=section_banner_width_cm,
        )
        elog("Inserted question difficulty tables: " + str(inserted_tables))

    if append_difficulty_report:
        report_data = difficulty_report_data or gerar_relatorio_dificuldade_por_secao(
            input_path,
            area_conhecimento=area_conhecimento,
        )
        if append_difficulty_report_appendix(
            doc,
            report_data,
            margin_top_cm=margin_top_cm,
            margin_bottom_cm=margin_bottom_cm,
            margin_left_cm=margin_left_cm,
            margin_right_cm=margin_right_cm,
        ):
            elog("Inserted A4 difficulty report appendix at document end.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    elog("Saved via python-docx OK")

    if finalize_word:
        from word_finalize import finalize_with_word

        finalize_with_word(
            docx_in=output_path,
            docx_out=output_path,
            column_width_cm=column_width_cm,
            font_name=font_name,
            font_size=font_size,
            justify=justify,
            force_inline_wrap=force_inline_wrap,
            badge_tag=BADGE_TAG,
        )
        elog("Finalized with Word OK")

    return {"arquivo_saida": str(output_path)}
